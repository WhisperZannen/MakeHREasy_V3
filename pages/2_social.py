# ==============================================================================
# 文件路径: pages/2_🛡️_社保与福利结算.py
# 功能描述: 多主体社保与福利结算中枢 (MVC 架构前端 UI 层 - 财务隔离终极版)
# 核心改造:
#   1. Tab 1: 完美分离 199 与 7。
#   2. Tab 2: 彻底修复 Grouper 1D 崩溃。实装 5 大独立财务提取单引擎（总览与月度拆分）。
#   3. Tab 2 第二部分：将零散的 ZIP 下载彻底替换为“一键生成全量合并 Word”。
#   4. Tab 4: 历史补录支持时间感知倒推（智能寻轨引擎）。
#   5. [终极封箱] 修复嵌套按钮导致的状态流失 Bug；补齐 Tab 3 参数配置的中文翻译。
# ==============================================================================

# ==============================================================================
# 文件路径: pages/2_social.py
# 功能描述: 多主体社保与福利结算中枢 (UI 体验终极重构版)
# ==============================================================================

import streamlit as st
import pandas as pd
import io
import uuid
import zipfile
import datetime
import json
import os

from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# 引入公文渲染引擎
try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    st.error("🚨 缺少 Word 生成引擎！请在终端运行：pip install python-docx")
    st.stop()

from modules.core_social_security import (
    get_policy_rules,
    upsert_policy_rules,
    _get_db_connection,
    batch_update_emp_matrix
)
from modules.core_arrangements import (
    ACTIVE_LABELS,
    AMOUNT_SOURCE_LABELS,
    ARRANGEMENT_LABELS,
    COST_BEARER_RULE_LABELS,
    ENABLED_LABELS,
    INSURANCE_LABELS,
    PAYER_RULE_LABELS,
    SETTLEMENT_BATCH_STATUS_LABELS,
    SETTLEMENT_CYCLE_LABELS,
    SETTLEMENT_MODE_LABELS,
    backfill_relationship_snapshots,
    create_route_policy,
    create_social_override,
    get_entities_dataframe,
    get_route_policies_dataframe,
    get_settlement_batches_dataframe,
    get_social_overrides_dataframe,
    register_social_settlement_batch,
    resolve_social_route,
    update_settlement_batch_status,
)

st.set_page_config(page_title="社保与福利结算", layout="wide")


# ==============================================================================
# [UI 增强] 财务级 Excel 自动排版引擎
# ==============================================================================
def format_excel_sheet(worksheet, df_columns):
    worksheet.freeze_panes = 'A2'  # 强制冻结首行
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'),
                         bottom=Side(style='thin'))
    header_fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")

    for col_idx, col_name in enumerate(df_columns, 1):
        col_letter = get_column_letter(col_idx)
        # 智能调节列宽：姓名工号窄点，金额和通道名宽点
        if col_name in ['姓名', '工号']:
            worksheet.column_dimensions[col_letter].width = 10
        elif '通道' in col_name or '主体' in col_name:
            worksheet.column_dimensions[col_letter].width = 18
        elif '月' in col_name:
            worksheet.column_dimensions[col_letter].width = 16
        else:
            worksheet.column_dimensions[col_letter].width = 14

        for row_idx in range(1, worksheet.max_row + 1):
            cell = worksheet[f"{col_letter}{row_idx}"]
            cell.border = thin_border
            if row_idx == 1:
                cell.font = Font(bold=True)
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal='center', vertical='center')
            else:
                # 凡是带有“金额”、“企”、“个”、“大病”、“统筹”、“合计”字眼的，数字靠右
                if any(x in col_name for x in ['金额', '企', '个', '大病', '统筹', '合计']):
                    cell.number_format = '#,##0.00'
                    cell.alignment = Alignment(horizontal='right', vertical='center')
                else:
                    cell.alignment = Alignment(horizontal='center', vertical='center')


# ==============================================================================
# [配置中枢] Settings.json
# ==============================================================================
def load_settings():
    current_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(current_dir)
    settings_path = os.path.join(project_root, 'settings.json')
    if not os.path.exists(settings_path):
        return {"bank_accounts": {}, "company_signature": "省公众人力部"}
    with open(settings_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def safe_float(val, default=0.0):
    try:
        if pd.notna(val) and val is not None and str(val).strip() != '': return float(val)
        return default
    except Exception:
        return default


if 'show_confirm' not in st.session_state: st.session_state['show_confirm'] = False
if 'pending_params' not in st.session_state: st.session_state['pending_params'] = None

st.title("🛡️ 社保与福利结算中心")
st.caption("核心业务流向：当月基数备料 ➡️ 理论核算与补缴对账 ➡️ 跨主体结算与公对公要款 ➡️ 引擎底座配置")

tab1, tab2, tab3, tab4, tab5 = st.tabs(
    ["🧮 当月社保沙盘 (含补缴)", "📤 财务提款与公对公结算函", "⚙️ 全局规则与参数配置",
     "📥 历史账单导入 (冷启动)", "🔀 社保险种缴费规则"])

# ------------------------------------------------------------------------------
# Tab 1: 当月社保沙盘与对账池
# ------------------------------------------------------------------------------
with tab1:
    st.info("💡 业务铁律：先在【第一步】确保所有人基数就绪，再在【第二步】跑出当期理论账单，最后在【第三步】补录官方滞纳金与历史补缴差额。")

    # ==========================================================================
    # [极致体验修复] 智能推算下个月份引擎
    # ==========================================================================
    conn = _get_db_connection()
    try:
        cur = conn.cursor()
        cur.execute("SELECT MAX(cost_month) FROM ss_monthly_records")
        max_m = cur.fetchone()[0]
        if max_m:
            y, m = map(int, max_m.split('-'))
            if m == 12:
                default_month = f"{y+1}-01"
            else:
                default_month = f"{y}-{m+1:02d}"
        else:
            default_month = datetime.date.today().strftime("%Y-%m")
    except Exception:
        default_month = datetime.date.today().strftime("%Y-%m")
    finally:
        conn.close()

    # 提示语明确告诉你要按回车
    calc_month = st.text_input("📅 当前核算工作月份 (修改后请按【回车键】确认👇)", value=default_month, max_chars=7)

    st.subheader("🛠️ 第一步：基数初始化与特例抢救")
    conn = _get_db_connection()
    try:
        # [极致修复 1] 强力复刻人员模块的排序算法：部门权重 -> 岗位权重 -> 岗级(负数取反)
        detect_sql = """
            SELECT 
                e.emp_id AS '工号', e.name AS '姓名', d.dept_name AS '部门', e.status AS '人事状态',
                IFNULL(m.cost_center, '本级') AS '财务归属', IFNULL(m.base_salary_avg, 0.0) AS '已录入原始基数',
                IFNULL(m.fund_base_avg, 0.0) AS '独立公积金基数(选填)',
                IFNULL(m.pension_enabled, 1) AS '养老参保(1是0否)', IFNULL(m.pension_account, '省公众') AS '养老缴纳主体',
                IFNULL(m.medical_enabled, 1) AS '医疗参保(1是0否)', IFNULL(m.medical_account, '省公司') AS '医疗缴纳主体',
                7.0 AS '大病统筹(个人固定)',
                IFNULL(m.unemp_enabled, 1) AS '失业参保(1是0否)', IFNULL(m.unemp_account, '省公众') AS '失业缴纳主体',
                IFNULL(m.injury_enabled, 1) AS '工伤参保(1是0否)', IFNULL(m.injury_account, '省公众') AS '工伤缴纳主体',
                IFNULL(m.maternity_enabled, 1) AS '生育参保(1是0否)', IFNULL(m.maternity_account, '省公司') AS '生育缴纳主体',
                IFNULL(m.fund_enabled, 1) AS '公积金参保(1是0否)', IFNULL(m.fund_account, '省公众') AS '公积金缴纳主体',
                IFNULL(m.annuity_enabled, 0) AS '年金参保(1是0否)', IFNULL(m.annuity_account, '省公司') AS '年金缴纳主体'
            FROM employees e
            LEFT JOIN departments d ON e.dept_id = d.dept_id
            LEFT JOIN employee_profiles p ON e.emp_id = p.emp_id
            LEFT JOIN positions pos ON p.pos_id = pos.pos_id
            LEFT JOIN ss_emp_matrix m ON e.emp_id = m.emp_id
            WHERE e.status IN ('在职', '挂靠人员')
            ORDER BY 
                CASE WHEN e.status = '退休' OR d.dept_name LIKE '%离退休%' THEN 9999
                     WHEN e.status = '挂靠人员' THEN 9000
                     WHEN d.dept_name LIKE '%公共%' OR d.dept_name LIKE '%统筹%' THEN 9998
                     ELSE IFNULL(d.sort_order, 999) END ASC,
                IFNULL(pos.sort_order, 999) ASC,
                CASE WHEN e.post_rank IS NOT NULL THEN -e.post_rank ELSE 9999.0 END ASC,
                e.emp_id ASC
        """
        roster_df = pd.read_sql_query(detect_sql, conn)
    finally:
        conn.close()

    c_down, c_up = st.columns(2)
    with c_down:
        # 基数表使用排版引擎冻结表头、拉宽列距
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            roster_df.to_excel(writer, index=False, sheet_name='全员基数与缴交主体配置表')
            format_excel_sheet(writer.sheets['全员基数与缴交主体配置表'], roster_df.columns)

        # [防呆设计] 按钮上直接大字显示当前认定的月份，绝对不会下错！
        st.download_button(
            label=f"📥 1. 下载【{calc_month}】基数初始化名单 (已排版)",
            data=buffer.getvalue(),
            file_name=f"全员基数表_{calc_month}.xlsx",
            mime="application/vnd.ms-excel"
        )

    with c_up:
        uploaded_file = st.file_uploader(f"📤 2. 上传填好的【{calc_month}】基数配置 Excel", type=["xlsx", "csv"], label_visibility="collapsed")
        if uploaded_file is not None:
            if st.session_state.get('last_processed_file_id') != uploaded_file.file_id:
                try:
                    upload_df = pd.read_csv(uploaded_file) if uploaded_file.name.endswith('.csv') else pd.read_excel(uploaded_file)
                    success, msg = batch_update_emp_matrix(upload_df)
                    if success:
                        st.session_state['last_processed_file_id'] = uploaded_file.file_id
                        st.rerun()
                    else: st.error(msg)
                except Exception as e: st.error(f"❌ 读取文件失败: {e}")
            else: st.success("✅ 基数已成功灌入引擎，可以开始核算！")

    st.divider()

    st.subheader("🧮 第二步：本月正常参保核算")
    rule_year_to_use = st.selectbox("⚙️ 选择本次套用的【规则年度】(如次年6月前沿用上年规则)",
                                    ["2024", "2025", "2026", "2027", "2028"], index=1)

    if st.button("🚀 启动引擎，生成当期理论账单", type="primary"):
        missing_count = len(roster_df[roster_df['已录入原始基数'] == 0.0])
        if missing_count > 0:
            st.error(f"🚨 警告：发现 {missing_count} 名员工基数为 0，请在第一步补充完整！")
        else:
            all_bills = []
            from modules.core_social_security import calculate_complete_bill

            for _, row in roster_df.iterrows():
                all_bills.append(calculate_complete_bill(row.to_dict(), rule_year_to_use, calc_month))
            st.session_state['temp_bills'] = pd.DataFrame(all_bills)

    if 'temp_bills' in st.session_state:
        raw_df_preview = st.session_state['temp_bills']
        export_df = raw_df_preview.copy()

        calculation_error_cols = [
            c for c in raw_df_preview.columns
            if c.startswith('__') and c.endswith('_calculation_error')
        ]
        calculation_errors = []
        for col in calculation_error_cols:
            calculation_errors.extend(
                raw_df_preview[col].dropna().astype(str).unique().tolist()
            )
        if calculation_errors:
            st.error(
                "🚨 检测到缺失的社保计算政策，相关险种已暂停计算，禁止直接固化："
                + "；".join(sorted(set(calculation_errors)))
            )

        cols_to_drop = ['injury_个', 'maternity_个'] + [
            c for c in export_df.columns if c.startswith('__')
        ]
        export_df = export_df.drop(columns=[c for c in cols_to_drop if c in export_df.columns])

        audit_rename_map = {
            'pension_企': '养老(企业)', 'pension_个': '养老(个人)', 'pension_route': '养老缴纳主体',
            'medical_企': '医疗(企业)', 'medical_个': '医疗(个人)', 'medical_route': '医疗缴纳主体',
            'medical_serious_个': '大病(个人)', 'medical_serious_pers': '大病(个人)', '大病_个': '大病(个人)',
            'unemp_企': '失业(企业)', 'unemp_个': '失业(个人)', 'unemp_route': '失业缴纳主体',
            'injury_企': '工伤(企业)', 'injury_route': '工伤缴纳主体',
            'maternity_企': '生育(企业)', 'maternity_route': '生育缴纳主体',
            'fund_企': '公积金(企业)', 'fund_个': '公积金(个人)', 'fund_route': '公积金缴纳主体',
            'annuity_企': '年金(企业)', 'annuity_个': '年金(个人)', 'annuity_route': '年金缴纳主体'
        }
        export_df = export_df.rename(columns=audit_rename_map)
        if '大病(个人)' not in export_df.columns: export_df['大病(个人)'] = 0.0

        ordered_front_cols = ['工号', '姓名', '财务归属', '合计企业缴纳', '合计个人扣款']
        detail_cols = [c for c in export_df.columns if c not in ordered_front_cols]
        export_df = export_df[ordered_front_cols + detail_cols]

        # 将生成的底表与 roster_df（已排好序的）进行 Merge，强行保证导出文件的行顺序与基数名单完全一致！
        sort_ref = roster_df[['工号']].copy()
        sort_ref['__order__'] = range(len(sort_ref))
        export_df = pd.merge(export_df, sort_ref, on='工号', how='left').sort_values('__order__').drop(
            columns=['__order__'])

        search_query = st.text_input("🔍 抽查指定员工 (输入姓名或工号进行过滤审核)", "")
        display_df = export_df
        if search_query:
            display_df = display_df[
                display_df['姓名'].str.contains(search_query, na=False) | display_df['工号'].str.contains(search_query,
                                                                                                          na=False)]
        st.dataframe(display_df[['工号', '姓名', '财务归属', '合计企业缴纳', '合计个人扣款']], use_container_width=True,
                     hide_index=True)

        st.write("---")
        c_audit, c_save = st.columns(2)

        with c_audit:
            # [极致修复 2] 彻底消灭群魔乱舞！在预览下载中，直接为你拆分出 5 张独立主体 Sheet！
            buffer_audit = io.BytesIO()
            with pd.ExcelWriter(buffer_audit, engine='openpyxl') as writer:
                # 1. 写入全量总表
                export_df.to_excel(writer, index=False, sheet_name='0.全量合并底稿')
                format_excel_sheet(writer.sheets['0.全量合并底稿'], export_df.columns)

                # 根据本月实际“付款通道+险种”动态拆分，主体切换后不再漏表。
                item_meta = {
                    'pension': '养老', 'medical': '医疗', 'unemp': '失业',
                    'injury': '工伤', 'maternity': '生育', 'fund': '公积金',
                    'annuity': '年金'
                }
                channel_groups = {}
                channel_lookup = {}
                for item_code, item_name in item_meta.items():
                    hidden_col = f'__{item_code}_payment_channel_code'
                    route_col = f'{item_code}_route'
                    if hidden_col not in raw_df_preview.columns or route_col not in raw_df_preview.columns:
                        continue
                    for _, source_row in raw_df_preview.iterrows():
                        route_name = str(source_row.get(route_col, '')).strip()
                        channel_code = str(source_row.get(hidden_col, '')).strip()
                        if route_name in {'', '不参保', 'nan'} or not channel_code:
                            continue
                        key = (channel_code, route_name)
                        channel_groups.setdefault(key, set()).add(item_name)
                        if item_code == 'medical':
                            channel_groups[key].add('大病')
                        channel_lookup[(str(source_row['工号']), item_name)] = channel_code
                        if item_code == 'medical':
                            channel_lookup[(str(source_row['工号']), '大病')] = channel_code

                split_configs = []
                for idx, ((channel_code, route_name), item_names) in enumerate(
                    sorted(channel_groups.items(), key=lambda value: (value[0][1], value[0][0])),
                    1
                ):
                    ordered_items = [
                        name for name in ['养老', '医疗', '大病', '失业', '工伤', '生育', '公积金', '年金']
                        if name in item_names
                    ]
                    short_items = '_'.join(ordered_items)
                    split_configs.append({
                        'name': f'{idx}.{route_name}({short_items})'[:31],
                        'route': route_name,
                        'channel': channel_code,
                        'items': ordered_items,
                    })

                # 2. 依次生成分表
                for cfg in split_configs:
                    df_sub = export_df.copy()
                    # 动态筛选出该主体关注的列名
                    cols_to_keep = ['工号', '姓名', '财务归属']

                    # 过滤逻辑：只有那些主体名字符合的，才保留金额；否则设为 0
                    has_money = pd.Series([False] * len(df_sub), index=df_sub.index)
                    for item in cfg['items']:
                        if item == '大病':
                            channel_mask = df_sub['工号'].astype(str).map(
                                lambda emp_id: channel_lookup.get((emp_id, '大病'))
                            ) == cfg['channel']
                            mask = (df_sub['医疗缴纳主体'] == cfg['route']) & channel_mask
                            df_sub.loc[~mask, '大病(个人)'] = 0.0
                            if '大病(个人)' in df_sub.columns: cols_to_keep.append('大病(个人)')
                            has_money = has_money | (df_sub['大病(个人)'] > 0)
                        else:
                            route_col = f"{item}缴纳主体"
                            channel_mask = df_sub['工号'].astype(str).map(
                                lambda emp_id, item_name=item: channel_lookup.get((emp_id, item_name))
                            ) == cfg['channel']
                            mask = (df_sub[route_col] == cfg['route']) & channel_mask

                            c_col = f"{item}(企业)"
                            p_col = f"{item}(个人)"
                            if c_col in df_sub.columns:
                                df_sub.loc[~mask, c_col] = 0.0
                                cols_to_keep.append(c_col)
                                has_money = has_money | (df_sub[c_col] > 0)
                            if p_col in df_sub.columns:
                                df_sub.loc[~mask, p_col] = 0.0
                                cols_to_keep.append(p_col)
                                has_money = has_money | (df_sub[p_col] > 0)

                    df_sub_clean = df_sub[has_money][cols_to_keep]
                    if not df_sub_clean.empty:
                        df_sub_clean.to_excel(writer, index=False, sheet_name=cfg['name'])
                        format_excel_sheet(writer.sheets[cfg['name']], df_sub_clean.columns)

            st.download_button(
                label="📥 1. 下载当月核算底稿 (全自动排版，包含 5 大拆分 Sheet)",
                data=buffer_audit.getvalue(),
                file_name=f"当期核算智能底稿_{calc_month}.xlsx",
                type="secondary"
            )

        with c_save:
            if st.button(
                "💾 2. 线下复核无误，将当期明细固化入库",
                type="primary",
                disabled=bool(calculation_errors),
            ):
                from modules.core_social_security import save_monthly_ss_records

                success, msg = save_monthly_ss_records(raw_df_preview, calc_month)
                if success:
                    st.success(msg)
                    del st.session_state['temp_bills']
                    st.rerun()
                else:
                    st.error(msg)

    st.subheader("📥 第三步：补缴与滞纳金手工入账 (对齐官方核定单)")
    st.write("🔧 遇到历史跨月补缴、滞纳金等系统无法自动推演的账目，请在此按社保局单据直接填报写入。")

    rc1, rc2 = st.columns(2)
    with rc1:
        # [极致修复 3] 将表头命名为极其详细的白痴级规范，让使用者一眼就懂格式
        retro_cols = [
            '处理月份(必填:YYYY-MM)', '工号(必填)', '补缴起始月(选填:YYYY-MM)', '补缴结束月(选填:YYYY-MM)',
            '补缴险种(必选下拉框)', '企业本金合计', '个人本金合计', '企业承担滞纳金', '备注(原因等)'
        ]
        retro_template = pd.DataFrame(columns=retro_cols)
        retro_buffer = io.BytesIO()
        with pd.ExcelWriter(retro_buffer, engine='openpyxl') as writer:
            retro_template.to_excel(writer, index=False, sheet_name='补缴规范模板')
            ws = writer.sheets['补缴规范模板']
            # 使用引擎排版并加宽列距
            format_excel_sheet(ws, retro_cols)
            # 对于日期要求的列特别标红提醒
            red_fill = PatternFill(start_color="FFCCCC", end_color="FFCCCC", fill_type="solid")
            for i in [1, 3, 4]:  # 处理月份, 起始月, 结束月
                ws.cell(row=1, column=i).fill = red_fill

            from openpyxl.worksheet.datavalidation import DataValidation

            dv = DataValidation(type="list",
                                formula1='"养老保险,医疗保险,大病医疗,失业保险,工伤保险,生育保险,住房公积金,企业年金"',
                                allow_blank=False)
            ws.add_data_validation(dv)
            dv.add("E2:E1000")  # 在E列加上下拉框限制

        st.download_button("📥 下载补缴与滞纳金专用模板 (已标明日历格式与下拉框)", data=retro_buffer.getvalue(),
                           file_name=f"补缴专用导入模板_{calc_month}.xlsx")

    with rc2:
        retro_file = st.file_uploader("📤 上传已填好的补缴核定单", type=["xlsx", "csv"], label_visibility="collapsed")
        if retro_file and st.button("💾 将补缴数据强行入库"):
            try:
                r_df = pd.read_csv(retro_file) if retro_file.name.endswith('.csv') else pd.read_excel(retro_file)
                conn = _get_db_connection()
                cursor = conn.cursor()
                sql = """
                      INSERT INTO ss_retroactive_records (retro_id, process_month, emp_id, target_start_month, \
                                                          target_end_month, retro_type, \
                                                          total_comp_retro, total_pers_retro, late_fee, remarks,
                                                          arrangement_id, payer_entity_code, cost_bearer_code,
                                                          settlement_counterparty_code, settlement_mode) \
                      VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?) \
                      """
                count = 0
                retro_item_map = {
                    '养老保险': 'pension', '医疗保险': 'medical', '大病医疗': 'medical',
                    '失业保险': 'unemp', '工伤保险': 'injury', '生育保险': 'maternity',
                    '住房公积金': 'fund', '企业年金': 'annuity'
                }
                matrix_account_columns = {
                    'pension': ('pension_enabled', 'pension_account'),
                    'medical': ('medical_enabled', 'medical_account'),
                    'unemp': ('unemp_enabled', 'unemp_account'),
                    'injury': ('injury_enabled', 'injury_account'),
                    'maternity': ('maternity_enabled', 'maternity_account'),
                    'fund': ('fund_enabled', 'fund_account'),
                    'annuity': ('annuity_enabled', 'annuity_account'),
                }
                for _, row in r_df.iterrows():
                    eid = str(row.get('工号(必填)', '')).replace('.0', '').strip()
                    if not eid or eid == 'nan': continue
                    process_month = str(row.get('处理月份(必填:YYYY-MM)', calc_month)).strip()
                    retro_type = str(row.get('补缴险种(必选下拉框)', '未知险种'))
                    item_code = retro_item_map.get(retro_type, 'pension')
                    matrix_row = cursor.execute(
                        "SELECT * FROM ss_emp_matrix WHERE emp_id = ?", (eid,)
                    ).fetchone()
                    matrix_data = dict(matrix_row) if matrix_row else {}
                    enabled_col, account_col = matrix_account_columns[item_code]
                    route_context = resolve_social_route(
                        eid, item_code, process_month,
                        legacy_enabled=int(matrix_data.get(enabled_col, 1) or 0),
                        legacy_payer_name=str(matrix_data.get(account_col) or '省公众'),
                        legacy_cost_center=str(matrix_data.get('cost_center') or '本级'),
                        conn=conn,
                    )
                    cursor.execute(sql, (
                        str(uuid.uuid4())[:12], process_month, eid,
                        str(row.get('补缴起始月(选填:YYYY-MM)', '')), str(row.get('补缴结束月(选填:YYYY-MM)', '')),
                        retro_type,
                        safe_float(row.get('企业本金合计', 0.0)), safe_float(row.get('个人本金合计', 0.0)),
                        safe_float(row.get('企业承担滞纳金', 0.0)), str(row.get('备注(原因等)', '')),
                        route_context.get('arrangement_id'), route_context.get('payer_entity_code'),
                        route_context.get('cost_bearer_code'), route_context.get('settlement_counterparty_code'),
                        route_context.get('settlement_mode') or 'none'
                    ))
                    count += 1
                conn.commit()
                st.success(f"✅ 成功将 {count} 笔特殊补缴与滞纳金记录封印入库！后续台账与工资计算将自动识别叠加。")
            except Exception as e:
                st.error(f"❌ 写入补缴表失败: {e}")
            finally:
                if 'conn' in locals(): conn.close()


# ------------------------------------------------------------------------------
# Tab 5: 带生效期的缴费路由与个人例外
# ------------------------------------------------------------------------------
with tab5:
    st.subheader("🔀 社保险种缴费规则与个人特殊处理")
    st.info(
        "系统按“单个人员特殊规则 ＞ 同类人员统一规则 ＞ 原参保配置”的顺序判断每个人、每个险种。"
        "这里只决定由谁计算、谁缴费、成本由谁承担以及是否结算，不直接修改已经封账的历史账单。"
    )

    entity_df = get_entities_dataframe(active_only=True)
    entity_names = dict(zip(entity_df['entity_code'], entity_df['entity_name']))
    entity_options = [None] + entity_df['entity_code'].tolist()
    item_options = list(INSURANCE_LABELS.keys())

    route_tab, override_tab = st.tabs([
        "📘 同类人员统一规则（公共政策）", "👤 单个人员特殊规则（个人例外）"
    ])

    with route_tab:
        st.caption(
            "用于一批人共同遵循的规则。例如：从某月开始，全部下沉人员的基本医疗由派出单位缴纳，"
            "费用年末与所在分公司结算。相同规则只需按险种配置一次。"
        )
        policy_df = get_route_policies_dataframe(active_only=False)
        if not policy_df.empty:
            display_policy = policy_df.copy()
            display_policy['关系类型'] = display_policy['arrangement_type'].map(ARRANGEMENT_LABELS).fillna(display_policy['arrangement_type'])
            display_policy['险种'] = display_policy['insurance_item'].map(INSURANCE_LABELS).fillna(display_policy['insurance_item'])
            display_policy['payer_entity_rule'] = display_policy['payer_entity_rule'].map(PAYER_RULE_LABELS).fillna(display_policy['payer_entity_rule'])
            display_policy['cost_bearer_rule'] = display_policy['cost_bearer_rule'].map(COST_BEARER_RULE_LABELS).fillna(display_policy['cost_bearer_rule'])
            display_policy['settlement_mode'] = display_policy['settlement_mode'].map(SETTLEMENT_MODE_LABELS).fillna(display_policy['settlement_mode'])
            display_policy['settlement_cycle'] = display_policy['settlement_cycle'].map(SETTLEMENT_CYCLE_LABELS).fillna(display_policy['settlement_cycle'])
            display_policy['amount_source'] = display_policy['amount_source'].map(AMOUNT_SOURCE_LABELS).fillna(display_policy['amount_source'])
            display_policy['active'] = display_policy['active'].map(ACTIVE_LABELS).fillna(display_policy['active'])
            st.dataframe(
                display_policy[[
                    'route_policy_id', 'policy_name', '关系类型', 'contract_entity_name', '险种',
                    'effective_from_month', 'effective_to_month', 'calculation_policy_entity_name',
                    'payer_entity_rule', 'payer_entity_name', 'cost_bearer_rule',
                    'cost_bearer_name', 'settlement_mode', 'settlement_cycle',
                    'amount_source', 'priority', 'active'
                ]].rename(columns={
                    'route_policy_id': '政策ID', 'policy_name': '政策名称',
                    'contract_entity_name': '适用牌子', 'effective_from_month': '生效月',
                    'effective_to_month': '失效月', 'calculation_policy_entity_name': '计算政策主体',
                    'payer_entity_rule': '缴费主体规则', 'payer_entity_name': '固定缴费主体',
                    'cost_bearer_rule': '成本承担规则', 'cost_bearer_name': '固定成本单位',
                    'settlement_mode': '结算方式', 'settlement_cycle': '结算周期',
                    'amount_source': '金额来源', 'priority': '优先级', 'active': '启用'
                }),
                use_container_width=True,
                hide_index=True,
            )
        else:
            st.caption("尚未设置公共路由政策，系统继续完全沿用现有人员参保矩阵。")

        with st.expander("➕ 新增一个生效版本", expanded=False):
            with st.form("route_policy_form"):
                policy_name = st.text_input("政策名称*")
                pc1, pc2, pc3 = st.columns(3)
                with pc1:
                    arrangement_type = st.selectbox(
                        "适用关系*", list(ARRANGEMENT_LABELS.keys()),
                        format_func=lambda value: ARRANGEMENT_LABELS[value]
                    )
                    contract_entity = st.selectbox(
                        "限定劳动合同主体", entity_options,
                        format_func=lambda value: "全部牌子" if value is None else entity_names[value]
                    )
                    insurance_item = st.selectbox(
                        "险种*", item_options,
                        format_func=lambda value: INSURANCE_LABELS[value]
                    )
                with pc2:
                    effective_from = st.text_input("生效月份*", value=datetime.date.today().strftime('%Y-%m'))
                    has_policy_end = st.checkbox("设置失效月份")
                    effective_to = st.text_input("失效月份", value="")
                    enabled_option = st.selectbox("参保开关", ['沿用个人配置', '默认参保', '默认不参保'])
                with pc3:
                    calculation_entity = st.selectbox(
                        "计算政策主体", entity_options,
                        format_func=lambda value: "沿用原缴费主体" if value is None else entity_names[value]
                    )
                    amount_source = st.selectbox(
                        "金额来源", list(AMOUNT_SOURCE_LABELS),
                        format_func=lambda value: AMOUNT_SOURCE_LABELS[value],
                    )
                    priority = st.number_input(
                        "优先级", min_value=1, max_value=9999, value=100,
                        help="同一人员、险种和月份匹配多条规则时，数字越大越优先。",
                    )

                pr1, pr2 = st.columns(2)
                with pr1:
                    payer_rule = st.selectbox(
                        "实际缴费主体规则",
                        list(PAYER_RULE_LABELS),
                        format_func=lambda value: PAYER_RULE_LABELS[value],
                    )
                    payer_entity = st.selectbox(
                        "固定缴费主体", entity_options,
                        format_func=lambda value: "不固定" if value is None else entity_names[value]
                    )
                    payment_channel = st.text_input("付款通道编码（可选）")
                with pr2:
                    cost_rule = st.selectbox(
                        "最终成本承担规则",
                        list(COST_BEARER_RULE_LABELS),
                        format_func=lambda value: COST_BEARER_RULE_LABELS[value],
                    )
                    cost_entity = st.selectbox(
                        "固定成本承担单位", entity_options,
                        format_func=lambda value: "不固定" if value is None else entity_names[value]
                    )
                    settlement_counterparty = st.selectbox(
                        "固定结算对象", entity_options,
                        format_func=lambda value: "按关系处理" if value is None else entity_names[value]
                    )

                ps1, ps2 = st.columns(2)
                with ps1:
                    settlement_mode = st.selectbox(
                        "结算方式", list(SETTLEMENT_MODE_LABELS),
                        format_func=lambda value: SETTLEMENT_MODE_LABELS[value],
                    )
                    settlement_cycle = st.selectbox(
                        "结算周期", list(SETTLEMENT_CYCLE_LABELS),
                        format_func=lambda value: SETTLEMENT_CYCLE_LABELS[value],
                    )
                with ps2:
                    policy_remarks = st.text_area("政策说明")

                if st.form_submit_button("保存新政策版本", type="primary"):
                    enabled_default = {'沿用个人配置': None, '默认参保': 1, '默认不参保': 0}[enabled_option]
                    ok, msg = create_route_policy({
                        'policy_name': policy_name,
                        'arrangement_type': arrangement_type,
                        'contract_entity_code': contract_entity,
                        'insurance_item': insurance_item,
                        'effective_from_month': effective_from.strip(),
                        'effective_to_month': effective_to.strip() if has_policy_end and effective_to.strip() else None,
                        'enabled_default': enabled_default,
                        'calculation_policy_entity': calculation_entity,
                        'payer_entity_rule': payer_rule,
                        'payer_entity_code': payer_entity,
                        'cost_bearer_rule': cost_rule,
                        'cost_bearer_code': cost_entity,
                        'settlement_counterparty_code': settlement_counterparty,
                        'settlement_mode': settlement_mode,
                        'settlement_cycle': settlement_cycle,
                        'amount_source': amount_source,
                        'payment_channel_code': payment_channel.strip() or None,
                        'priority': int(priority), 'active': 1, 'remarks': policy_remarks,
                    })
                    if ok:
                        st.success(msg)
                    else:
                        st.error(msg)

    with override_tab:
        st.caption(
            "用于只针对某个人、某个险种的例外，优先级最高。例如李峰林只有工伤由省公众缴纳，"
            "就只新增一条“李峰林－工伤”例外，不影响他的其他险种，也不影响其他人。"
        )
        override_df = get_social_overrides_dataframe(active_only=False)
        if not override_df.empty:
            display_override = override_df.copy()
            display_override['险种'] = display_override['insurance_item'].map(INSURANCE_LABELS).fillna(display_override['insurance_item'])
            display_override['enabled'] = display_override['enabled'].map(ENABLED_LABELS).fillna(display_override['enabled'])
            display_override['settlement_mode'] = display_override['settlement_mode'].map(SETTLEMENT_MODE_LABELS).fillna(display_override['settlement_mode'])
            display_override['settlement_cycle'] = display_override['settlement_cycle'].map(SETTLEMENT_CYCLE_LABELS).fillna(display_override['settlement_cycle'])
            display_override['amount_source'] = display_override['amount_source'].map(AMOUNT_SOURCE_LABELS).fillna(display_override['amount_source'])
            display_override['active'] = display_override['active'].map(ACTIVE_LABELS).fillna(display_override['active'])
            st.dataframe(
                display_override[[
                    'override_id', 'emp_name', 'emp_id', '险种', 'effective_from_month',
                    'effective_to_month', 'enabled', 'calculation_policy_entity_name',
                    'payer_entity_name', 'cost_bearer_name', 'settlement_counterparty_name',
                    'settlement_mode', 'settlement_cycle', 'amount_source',
                    'special_reason', 'source_document_no', 'active'
                ]].rename(columns={
                    'override_id': '例外ID', 'emp_name': '姓名', 'emp_id': '工号',
                    'effective_from_month': '生效月', 'effective_to_month': '失效月',
                    'enabled': '参保', 'calculation_policy_entity_name': '计算政策主体',
                    'payer_entity_name': '缴费主体', 'cost_bearer_name': '成本承担',
                    'settlement_counterparty_name': '结算对象',
                    'settlement_mode': '结算方式', 'settlement_cycle': '结算周期',
                    'amount_source': '金额来源',
                    'special_reason': '特殊原因', 'source_document_no': '依据', 'active': '启用'
                }),
                use_container_width=True,
                hide_index=True,
            )
        else:
            st.caption("暂无个人例外。原挂靠人员仍按现有Excel参保开关和缴费主体运行。")

        conn_people = _get_db_connection()
        people_df = pd.read_sql_query(
            "SELECT emp_id, name, status FROM employees WHERE status IN ('在职','挂靠人员') ORDER BY name",
            conn_people
        )
        conn_people.close()
        people_options = people_df['emp_id'].astype(str).tolist()
        people_names = dict(zip(
            people_df['emp_id'].astype(str),
            people_df.apply(lambda row: f"{row['name']}（{row['status']}）", axis=1)
        ))

        with st.expander("➕ 新增个人险种例外", expanded=False):
            with st.form("social_override_form"):
                oc1, oc2, oc3 = st.columns(3)
                with oc1:
                    override_emp = st.selectbox("员工*", people_options, format_func=lambda value: people_names[value])
                    override_item = st.selectbox("险种*", item_options, format_func=lambda value: INSURANCE_LABELS[value])
                    override_enabled = st.selectbox("参保开关", ['沿用原配置', '参保', '不参保'])
                with oc2:
                    override_from = st.text_input("生效月份*", value=datetime.date.today().strftime('%Y-%m'), key='override_from')
                    override_has_end = st.checkbox("设置失效月份", key='override_has_end')
                    override_to = st.text_input("失效月份", value="", key='override_to')
                    override_calc_entity = st.selectbox(
                        "计算政策主体", entity_options,
                        format_func=lambda value: "沿用公共/原配置" if value is None else entity_names[value],
                        key='override_calc_entity'
                    )
                with oc3:
                    override_payer = st.selectbox(
                        "实际缴费主体", entity_options,
                        format_func=lambda value: "沿用公共/原配置" if value is None else entity_names[value],
                        key='override_payer'
                    )
                    override_cost = st.selectbox(
                        "最终成本承担", entity_options,
                        format_func=lambda value: "沿用关系/原配置" if value is None else entity_names[value],
                        key='override_cost'
                    )
                    override_counterparty = st.selectbox(
                        "结算对象", entity_options,
                        format_func=lambda value: "沿用关系" if value is None else entity_names[value],
                        key='override_counterparty'
                    )

                oo1, oo2 = st.columns(2)
                with oo1:
                    override_settlement = st.selectbox(
                        "结算方式", [None] + list(SETTLEMENT_MODE_LABELS),
                        format_func=lambda value: "沿用人员关系/公共政策" if value is None else SETTLEMENT_MODE_LABELS[value],
                    )
                    override_cycle = st.selectbox(
                        "结算周期", [None] + list(SETTLEMENT_CYCLE_LABELS),
                        format_func=lambda value: "沿用人员关系/公共政策" if value is None else SETTLEMENT_CYCLE_LABELS[value],
                    )
                    override_amount_source = st.selectbox(
                        "金额来源", [None] + list(AMOUNT_SOURCE_LABELS),
                        format_func=lambda value: "沿用公共政策" if value is None else AMOUNT_SOURCE_LABELS[value],
                    )
                with oo2:
                    special_reason = st.text_input("特殊原因*")
                    override_document = st.text_input("政策/协议依据")
                    override_channel = st.text_input("付款通道编码（可选）")

                if st.form_submit_button("保存个人例外", type="primary"):
                    enabled_value = {'沿用原配置': None, '参保': 1, '不参保': 0}[override_enabled]
                    ok, msg = create_social_override({
                        'emp_id': override_emp, 'insurance_item': override_item,
                        'effective_from_month': override_from.strip(),
                        'effective_to_month': override_to.strip() if override_has_end and override_to.strip() else None,
                        'enabled': enabled_value,
                        'calculation_policy_entity': override_calc_entity,
                        'payer_entity_code': override_payer,
                        'cost_bearer_code': override_cost,
                        'settlement_counterparty_code': override_counterparty,
                        'settlement_mode': override_settlement,
                        'settlement_cycle': override_cycle,
                        'amount_source': override_amount_source,
                        'payment_channel_code': override_channel.strip() or None,
                        'special_reason': special_reason,
                        'source_document_no': override_document,
                        'active': 1,
                    })
                    if ok:
                        st.success(msg)
                    else:
                        st.error(msg)

# ------------------------------------------------------------------------------
# (Tab 2, Tab 3, Tab 4 维持之前的代码逻辑不动)
# ... 这里不再重复展示那些不需要修改的 Tab 2/3/4 ...
# 如果你需要，请确保将你之前的 Tab 2、3、4 的代码拼接到这里下面。
# ------------------------------------------------------------------------------

# ------------------------------------------------------------------------------
# Tab 2: 财务输出中心 (对内账单与公对公 Word 结算函)
# ------------------------------------------------------------------------------
with tab2:
    st.subheader("📤 第一部分：对内审批提款单 (跨期多表智能打包)")
    st.info("💡 财务走账专用。严格按【缴费主体+险种】物理隔离出 5 个独立的 Excel。多月提取时自动生成【总览】与月度明细 Sheet。")

    conn = _get_db_connection()
    try:
        available_months = pd.read_sql_query("SELECT DISTINCT cost_month FROM ss_monthly_records ORDER BY cost_month DESC", conn)['cost_month'].tolist()
    except Exception:
        available_months = []
    finally:
        conn.close()

    ic1, ic2 = st.columns(2)
    with ic1:
        int_start_month = st.selectbox("📅 对内请款起始月", options=available_months if available_months else ["无数据"], key='int_start', index=len(available_months)-1 if available_months else 0)
    with ic2:
        int_end_month = st.selectbox("📅 对内请款结束月", options=available_months if available_months else ["无数据"], key='int_end', index=0 if available_months else 0)

    # [修复点：防跳跃] 生成与下载分离
    if st.button("🚀 1. 极速分析并打包对内提款单 (ZIP)", type="primary") and int_start_month != "无数据":
        s_m, e_m = min(int_start_month, int_end_month), max(int_start_month, int_end_month)
        selected_months = [m for m in available_months if s_m <= m <= e_m]
        selected_months.sort()

        conn = _get_db_connection()

        query = """
            SELECT r.*, e.name AS '姓名'
            FROM ss_monthly_records r 
            LEFT JOIN employees e ON r.emp_id = e.emp_id 
            WHERE r.cost_month >= ? AND r.cost_month <= ?
        """
        raw_df = pd.read_sql_query(query, conn, params=[s_m, e_m])

        retro_query = """
            SELECT r.*, e.name AS '姓名',
                   COALESCE(cost_entity.entity_name, m.cost_center, '本级') AS 'cost_center',
                   payer_entity.entity_name AS 'retro_payer_name'
            FROM ss_retroactive_records r
            LEFT JOIN employees e ON r.emp_id = e.emp_id
            LEFT JOIN ss_emp_matrix m ON r.emp_id = m.emp_id
            LEFT JOIN business_entities cost_entity ON r.cost_bearer_code = cost_entity.entity_code
            LEFT JOIN business_entities payer_entity ON r.payer_entity_code = payer_entity.entity_code
            WHERE r.process_month >= ? AND r.process_month <= ?
        """
        retro_df = pd.read_sql_query(retro_query, conn, params=[s_m, e_m])
        conn.close()

        if not raw_df.empty or not retro_df.empty:
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:

                rename_map = {
                    'emp_id': '工号', 'cost_center': '财务归属',
                    'pension_comp': '养老(企业)', 'pension_pers': '养老(个人)',
                    'medical_comp': '医疗(企业)', 'medical_pers': '医疗(个人)', 'medical_serious_pers': '大病(个人)',
                    'unemp_comp': '失业(企业)', 'unemp_pers': '失业(个人)',
                    'injury_comp': '工伤(企业)', 'maternity_comp': '生育(企业)',
                    'fund_comp': '公积金(企业)', 'fund_pers': '公积金(个人)',
                    'annuity_comp': '年金(企业)', 'annuity_pers': '年金(个人)'
                }

                channel_configs = [
                    {'name': '1.中电数智(五险两金综合)', 'route': '中电数智', 'items': ['pension', 'medical', 'unemp', 'injury', 'maternity', 'fund', 'annuity']},
                    {'name': '2.省公司(年金)', 'route': '省公司', 'items': ['annuity']},
                    {'name': '3.省公司(医疗_生育_工伤)', 'route': '省公司', 'items': ['medical', 'maternity', 'injury']},
                    {'name': '4.省公众(公积金)', 'route': '省公众', 'items': ['fund']},
                    {'name': '5.省公众(养老_失业_工伤)', 'route': '省公众', 'items': ['pension', 'unemp', 'injury']}
                ]

                all_insurance_items = ['pension', 'medical', 'unemp', 'injury', 'maternity', 'fund', 'annuity']

                # 保留现有5个财务通道，同时自动发现主体切换后新增的“主体+险种”组合。
                # 例如医疗、生育转到省公众后，会形成新增通道文件，不会被旧硬编码漏掉。
                covered_pairs = {
                    (cfg['route'], item)
                    for cfg in channel_configs
                    for item in cfg['items']
                }
                extra_pairs = {}
                if not raw_df.empty:
                    for item in all_insurance_items:
                        route_col = f'{item}_route'
                        money_cols_for_item = [
                            col for col in [f'{item}_comp', f'{item}_pers']
                            if col in raw_df.columns
                        ]
                        if route_col not in raw_df.columns or not money_cols_for_item:
                            continue
                        positive_rows = raw_df[raw_df[money_cols_for_item].fillna(0).abs().sum(axis=1) > 0]
                        for route_name in positive_rows[route_col].dropna().astype(str).unique():
                            if route_name not in {'', '不参保', 'None'} and (route_name, item) not in covered_pairs:
                                extra_pairs.setdefault(route_name, set()).add(item)

                for route_name, items in sorted(extra_pairs.items()):
                    channel_configs.append({
                        'name': f"{len(channel_configs) + 1}.{route_name}(新增缴费通道)",
                        'route': route_name,
                        'items': sorted(items, key=all_insurance_items.index)
                    })

                if not raw_df.empty:
                    for config in channel_configs:
                        df_channel = raw_df.copy()

                        for it in all_insurance_items:
                            if it not in config['items']:
                                if f'{it}_comp' in df_channel.columns: df_channel[f'{it}_comp'] = 0.0
                                if f'{it}_pers' in df_channel.columns: df_channel[f'{it}_pers'] = 0.0
                                if it == 'medical' and 'medical_serious_pers' in df_channel.columns: df_channel['medical_serious_pers'] = 0.0
                            else:
                                mask = df_channel[f'{it}_route'] == config['route']
                                if f'{it}_comp' in df_channel.columns: df_channel.loc[~mask, f'{it}_comp'] = 0.0
                                if f'{it}_pers' in df_channel.columns: df_channel.loc[~mask, f'{it}_pers'] = 0.0
                                if it == 'medical' and 'medical_serious_pers' in df_channel.columns: df_channel.loc[~mask, 'medical_serious_pers'] = 0.0

                        money_cols = [c for c in df_channel.columns if c.endswith('_comp') or c.endswith('_pers')]
                        df_channel['__row_sum__'] = df_channel[money_cols].sum(axis=1)
                        df_channel = df_channel[df_channel['__row_sum__'] > 0].drop(columns=['__row_sum__'])

                        if not df_channel.empty:
                            excel_io = io.BytesIO()
                            with pd.ExcelWriter(excel_io, engine='xlsxwriter') as writer:
                                group_cols = ['emp_id', '姓名', 'cost_center']
                                active_sum_cols = [c for c in money_cols if df_channel[c].sum() > 0]

                                if active_sum_cols:
                                    df_sum = df_channel.groupby(group_cols)[active_sum_cols].sum().reset_index()
                                    export_cols = group_cols + active_sum_cols
                                    df_export = df_sum[export_cols].rename(columns=rename_map)
                                    df_export.to_excel(writer, index=False, sheet_name="总览")

                                for month in selected_months:
                                    df_month = df_channel[df_channel['cost_month'] == month]
                                    if not df_month.empty:
                                        m_active_cols = [c for c in money_cols if df_month[c].sum() > 0]
                                        if m_active_cols:
                                            export_cols = group_cols + m_active_cols
                                            df_export = df_month[export_cols].rename(columns=rename_map)
                                            df_export.to_excel(writer, index=False, sheet_name=month)

                            zf.writestr(f"{config['name']}_{s_m}至{e_m}.xlsx", excel_io.getvalue())

                if not retro_df.empty:
                    retro_map = {
                        'process_month': '处理月份', 'emp_id': '工号', 'retro_type': '补缴险种',
                        'target_start_month': '补缴起始月', 'target_end_month': '补缴结束月',
                        'total_comp_retro': '企业本金', 'total_pers_retro': '个人本金',
                        'late_fee': '滞纳金(异常支出)', 'remarks': '产生原因(备注)', 'cost_center': '财务归属'
                    }
                    retro_cols = ['处理月份', '工号', '姓名', '财务归属', '补缴险种', '补缴起始月', '补缴结束月', '企业本金', '个人本金', '滞纳金(异常支出)', '产生原因(备注)']
                    df_retro_export = retro_df.rename(columns=retro_map)[retro_cols]

                    excel_io = io.BytesIO()
                    with pd.ExcelWriter(excel_io, engine='xlsxwriter') as writer:
                        df_retro_export.to_excel(writer, index=False, sheet_name="异常款项专项审批")
                        workbook = writer.book
                        worksheet = writer.sheets['异常款项专项审批']
                        alert_format = workbook.add_format({'bg_color': '#FFC7CE', 'font_color': '#9C0006', 'bold': True})
                        worksheet.set_column('J:J', 16, alert_format)
                        worksheet.set_column('K:K', 35)

                    zf.writestr(f"6.异常款项专项审批_{s_m}至{e_m}.xlsx", excel_io.getvalue())

            # 将 ZIP 存入缓存记忆
            st.session_state['ss_zip_data'] = zip_buffer.getvalue()
            st.session_state['ss_zip_filename'] = f"对内审计提款单合集_{s_m}至{e_m}.zip"
            st.success("✅ 对内提款单已分析打包完毕，请点击下方按钮安全下载！")

    # 外置的下载按钮，随便点都不会跳回 Tab1
    if 'ss_zip_data' in st.session_state:
        st.download_button(
            label="📥 2. 点击下载对内提款单大包 (ZIP)",
            data=st.session_state['ss_zip_data'],
            file_name=st.session_state['ss_zip_filename'],
            type="secondary"
        )

    st.divider()

    # ==========================================
    # 第二部分：对外公对公结算函生成引擎 (一键合并打印版)
    # ==========================================
    st.subheader("📜 第二部分：跨期公对公结算函 (Word 动态打包引擎)")
    st.write("🔧 动态侦测全险种特例，自适应伸缩表格。彻底融合【跨月补缴与滞纳金】，精准追溯每一分财务死账。")

    ec1, ec2 = st.columns(2)
    with ec1: start_month = st.selectbox("⏳ 结算起始月", options=available_months if available_months else ["无数据"], key='s_month', index=len(available_months)-1 if available_months else 0)
    with ec2: end_month = st.selectbox("⏳ 结算结束月", options=available_months if available_months else ["无数据"], key='e_month', index=0 if available_months else 0)

    conn = _get_db_connection()
    s_m, e_m = min(start_month, end_month), max(start_month, end_month)
    branch_query = "SELECT DISTINCT cost_center FROM ss_monthly_records WHERE cost_month >= ? AND cost_month <= ? AND cost_center != '本级'"
    avail_branches_df = pd.read_sql_query(branch_query, conn, params=[s_m, e_m])
    conn.close()
    avail_branches = avail_branches_df['cost_center'].dropna().tolist()

    selected_branches = st.multiselect("🏢 勾选需要生成结算函的地市分公司 (默认全选)", options=avail_branches, default=avail_branches)

    # [修复点：防跳跃] 生成与下载分离
    if st.button("🚀 1. 一键排版【合并版】地市结算函 (全量 Word)", type="primary") and start_month != "无数据" and selected_branches:
        sys_settings = load_settings()
        conn = _get_db_connection()
        placeholders = ",".join(["?"] * len(selected_branches))
        ext_params = [s_m, e_m] + selected_branches

        ext_query = f"""
            SELECT r.*, e.name AS '姓名', 
                   COALESCE(NULLIF(m.ss_base_actual, 0.0), NULLIF(m.base_salary_avg, 0.0), 0.0) AS '基数'
            FROM ss_monthly_records r 
            LEFT JOIN employees e ON r.emp_id = e.emp_id 
            LEFT JOIN ss_emp_matrix m ON r.emp_id = m.emp_id
            WHERE r.cost_month >= ? AND r.cost_month <= ? AND r.cost_center IN ({placeholders})
        """
        ext_df = pd.read_sql_query(ext_query, conn, params=ext_params)

        item_mode_df = pd.read_sql_query(
            """
            SELECT cost_month, emp_id, insurance_item, settlement_mode
            FROM social_monthly_items
            WHERE cost_month >= ? AND cost_month <= ?
            """,
            conn,
            params=[s_m, e_m]
        )
        if not item_mode_df.empty and not ext_df.empty:
            settlement_mode_map = {
                (str(row['cost_month']), str(row['emp_id']), str(row['insurance_item'])):
                    str(row['settlement_mode'])
                for _, row in item_mode_df.iterrows()
            }
            allowed_letter_modes = {
                'proxy_social', 'annual_reimbursement', 'central_chargeback'
            }
            for idx, ext_row in ext_df.iterrows():
                month_key = str(ext_row['cost_month'])
                emp_key = str(ext_row['emp_id'])
                for item in ['pension', 'medical', 'unemp', 'injury', 'maternity', 'fund', 'annuity']:
                    mode = settlement_mode_map.get((month_key, emp_key, item))
                    if mode is not None and mode not in allowed_letter_modes:
                        if f'{item}_comp' in ext_df.columns:
                            ext_df.at[idx, f'{item}_comp'] = 0.0
                        if f'{item}_pers' in ext_df.columns:
                            ext_df.at[idx, f'{item}_pers'] = 0.0
                        if item == 'medical' and 'medical_serious_pers' in ext_df.columns:
                            ext_df.at[idx, 'medical_serious_pers'] = 0.0

        retro_query = f"""
            SELECT r.*, e.name AS '姓名', 
                   COALESCE(cost_entity.entity_name, m.cost_center, '本级') AS 'cost_center',
                   COALESCE(NULLIF(m.ss_base_actual, 0.0), NULLIF(m.base_salary_avg, 0.0), 0.0) AS '基数',
                   payer_entity.entity_name AS 'retro_payer_name',
                   m.pension_account AS 'pension_route', m.medical_account AS 'medical_route',
                   m.unemp_account AS 'unemp_route', m.injury_account AS 'injury_route',
                   m.maternity_account AS 'maternity_route', m.fund_account AS 'fund_route',
                   m.annuity_account AS 'annuity_route'
            FROM ss_retroactive_records r
            LEFT JOIN employees e ON r.emp_id = e.emp_id
            LEFT JOIN ss_emp_matrix m ON r.emp_id = m.emp_id
            LEFT JOIN business_entities cost_entity ON r.cost_bearer_code = cost_entity.entity_code
            LEFT JOIN business_entities payer_entity ON r.payer_entity_code = payer_entity.entity_code
            WHERE r.process_month >= ? AND r.process_month <= ?
              AND COALESCE(cost_entity.entity_name, m.cost_center, '本级') IN ({placeholders})
        """
        retro_df = pd.read_sql_query(retro_query, conn, params=ext_params)
        conn.close()

        retro_map = {'养老保险':'pension', '医疗保险':'medical', '大病医疗':'medical_serious', '失业保险':'unemp', '工伤保险':'injury', '生育保险':'maternity', '住房公积金':'fund', '企业年金':'annuity'}
        normalized_retro = []
        for _, row in retro_df.iterrows():
            retro_settlement_mode = row.get('settlement_mode')
            if pd.notna(retro_settlement_mode) and str(retro_settlement_mode) not in {
                '', 'proxy_social', 'annual_reimbursement', 'central_chargeback'
            }:
                continue
            prefix = retro_map.get(row.get('retro_type', ''))
            if not prefix: continue

            new_row = {
                'cost_month': f"补缴({row['process_month']})",
                'emp_id': row['emp_id'], '姓名': row['姓名'], 'cost_center': row['cost_center'], '基数': row['基数'],
                f'{prefix}_comp': row['total_comp_retro'], f'{prefix}_pers': row['total_pers_retro'],
                f'{prefix}_route': row.get('retro_payer_name') or row.get(f'{prefix}_route', '未知'),
                'late_fee': row['late_fee']
            }
            normalized_retro.append(new_row)

        retro_normalized_df = pd.DataFrame(normalized_retro)
        if not retro_normalized_df.empty:
            combined_df = pd.concat([ext_df, retro_normalized_df], ignore_index=True).fillna(0.0)
        else:
            combined_df = ext_df.copy()
            combined_df['late_fee'] = 0.0

        if combined_df.empty:
            st.warning("该时间段内所选分公司没有产生任何费用记录。")
        else:
            merged_doc = Document()
            first_letter = True

            for cc, group in combined_df.groupby('cost_center'):
                routes_used = set()
                for r_col in ['pension_route', 'medical_route', 'unemp_route', 'injury_route', 'maternity_route', 'fund_route', 'annuity_route']:
                    if r_col in group.columns: routes_used.update(group[r_col].astype(str).dropna().unique())
                routes_used.discard(''); routes_used.discard('None'); routes_used.discard('0.0')

                for route_name in routes_used:
                    df_cc_route = group.copy()
                    has_amt = pd.Series([False] * len(df_cc_route))

                    for it in ['pension', 'medical', 'unemp', 'injury', 'maternity', 'fund', 'annuity']:
                        r_col = f'{it}_route'
                        if r_col in df_cc_route.columns:
                            mask = df_cc_route[r_col].astype(str) == route_name
                            df_cc_route.loc[~mask, f'{it}_comp'] = 0.0
                            df_cc_route.loc[~mask, f'{it}_pers'] = 0.0
                            if it == 'medical' and 'medical_serious_pers' in df_cc_route.columns:
                                df_cc_route.loc[~mask, 'medical_serious_pers'] = 0.0

                            c_col_val = df_cc_route[f'{it}_comp'] if f'{it}_comp' in df_cc_route.columns else 0.0
                            p_col_val = df_cc_route[f'{it}_pers'] if f'{it}_pers' in df_cc_route.columns else 0.0
                            has_amt = has_amt | (c_col_val > 0) | (p_col_val > 0)

                    if 'late_fee' in df_cc_route.columns:
                        df_cc_route.loc[~has_amt, 'late_fee'] = 0.0
                        has_amt = has_amt | (df_cc_route['late_fee'] > 0)

                    df_cc_route = df_cc_route[has_amt]
                    if df_cc_route.empty: continue

                    active_cols = []
                    ins_names = {'pension':'养老', 'medical':'医疗', 'medical_serious':'大病', 'unemp':'失业', 'injury':'工伤', 'maternity':'生育', 'fund':'公积金', 'annuity':'年金'}

                    for it in ['pension', 'medical', 'unemp', 'injury', 'maternity', 'fund', 'annuity']:
                        c_sum = df_cc_route[f'{it}_comp'].sum() if f'{it}_comp' in df_cc_route.columns else 0
                        p_sum = df_cc_route[f'{it}_pers'].sum() if f'{it}_pers' in df_cc_route.columns else 0
                        if c_sum > 0 and p_sum > 0: active_cols.append({'id':it, 'name':ins_names[it], 'has_c':True, 'has_p':True})
                        elif c_sum > 0: active_cols.append({'id':it, 'name':ins_names[it], 'has_c':True, 'has_p':False})
                        elif p_sum > 0: active_cols.append({'id':it, 'name':ins_names[it], 'has_c':False, 'has_p':True})

                        if it == 'medical' and 'medical_serious_pers' in df_cc_route.columns:
                            if df_cc_route['medical_serious_pers'].sum() > 0:
                                active_cols.append({'id':'medical_serious', 'name':'大病', 'has_c':False, 'has_p':True})

                    has_late_fee = df_cc_route['late_fee'].sum() > 0 if 'late_fee' in df_cc_route.columns else False

                    row_totals = []
                    for _, r_data in df_cc_route.iterrows():
                        r_tot = 0.0
                        for ac in active_cols:
                            if ac['has_c']: r_tot += r_data.get(f"{ac['id']}_comp", 0.0)
                            if ac['has_p']: r_tot += r_data.get(f"{ac['id']}_pers", 0.0)
                        if has_late_fee: r_tot += r_data.get('late_fee', 0.0)
                        row_totals.append(r_tot)
                    df_cc_route['当行合计'] = row_totals
                    total_sum = sum(row_totals)
                    settlement_batch_id = register_social_settlement_batch(
                        s_m, e_m, str(cc), str(route_name), total_sum
                    )

                    if not first_letter: merged_doc.add_page_break()
                    first_letter = False

                    p_title = merged_doc.add_paragraph()
                    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r_title = p_title.add_run("关于社保代缴的结算说明函")
                    r_title.font.size = Pt(16); r_title.font.bold = True; r_title.font.name = '黑体'
                    r_title._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

                    p_salut = merged_doc.add_paragraph()
                    r_salut = p_salut.add_run(f"{cc}：")
                    r_salut.font.name = '宋体'; r_salut._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                    emp_names = "、".join(df_cc_route['姓名'].unique())
                    if len(emp_names) > 15: emp_names = emp_names[:15] + "等"

                    bank_info = sys_settings.get("bank_accounts", {}).get(route_name, {})
                    account_name = bank_info.get("名称", route_name)

                    ins_str = "、".join([ac['name'] for ac in active_cols])
                    p_body1 = merged_doc.add_paragraph(f"    因业务开展需要，{cc}{emp_names}社保（{ins_str}）暂由{account_name}代缴，代缴金额据实结算。")
                    p_body2 = merged_doc.add_paragraph(
                        f"    从{s_m[:4]}年{s_m[-2:]}月到{e_m[:4]}年{e_m[-2:]}月，"
                        f"代缴金额为{total_sum:.2f}元，明细如下：\n"
                        f"    系统结算批次：{settlement_batch_id}"
                    )
                    for p in [p_body1, p_body2]:
                        for run in p.runs:
                            run.font.name = '宋体'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体'); run.font.size = Pt(12)

                    base_headers = ["月份", "公司", "员工姓名", "基数"]
                    num_cols = len(base_headers)
                    for ac in active_cols:
                        if ac['has_c'] and ac['has_p']: num_cols += 2
                        else: num_cols += 1
                    if has_late_fee: num_cols += 1
                    num_cols += 1

                    table = merged_doc.add_table(rows=2 + len(df_cc_route) + 1, cols=num_cols)
                    table.style = 'Table Grid'

                    for i, h in enumerate(base_headers):
                        table.cell(0, i).merge(table.cell(1, i)).text = h

                    col_idx = len(base_headers)
                    for ac in active_cols:
                        if ac['has_c'] and ac['has_p']:
                            table.cell(0, col_idx).merge(table.cell(0, col_idx+1)).text = ac['name']
                            table.cell(1, col_idx).text = "企业"
                            table.cell(1, col_idx+1).text = "个人"
                            ac['c_idx'] = col_idx; ac['p_idx'] = col_idx + 1
                            col_idx += 2
                        elif ac['has_c']:
                            table.cell(0, col_idx).merge(table.cell(1, col_idx)).text = ac['name'] + "(企)"
                            ac['c_idx'] = col_idx; ac['p_idx'] = -1
                            col_idx += 1
                        elif ac['has_p']:
                            table.cell(0, col_idx).merge(table.cell(1, col_idx)).text = ac['name'] + "(个)"
                            ac['c_idx'] = -1; ac['p_idx'] = col_idx
                            col_idx += 1

                    if has_late_fee:
                        table.cell(0, col_idx).merge(table.cell(1, col_idx)).text = "滞纳金"
                        lf_idx = col_idx
                        col_idx += 1

                    table.cell(0, col_idx).merge(table.cell(1, col_idx)).text = "合计"
                    total_idx = col_idx

                    row_idx = 2
                    for _, r_data in df_cc_route.iterrows():
                        month_str = str(r_data['cost_month']).replace('补缴(', '').replace(')', '')
                        if '-' in month_str: month_str = month_str.split('-')[-1] + "月"

                        table.cell(row_idx, 0).text = month_str
                        table.cell(row_idx, 1).text = cc
                        table.cell(row_idx, 2).text = str(r_data['姓名'])
                        table.cell(row_idx, 3).text = str(r_data['基数'])

                        for ac in active_cols:
                            if ac['has_c']: table.cell(row_idx, ac['c_idx']).text = f"{r_data.get(ac['id']+'_comp', 0.0):.2f}"
                            if ac['has_p']: table.cell(row_idx, ac['p_idx']).text = f"{r_data.get(ac['id']+'_pers', 0.0):.2f}"

                        if has_late_fee: table.cell(row_idx, lf_idx).text = f"{r_data.get('late_fee', 0.0):.2f}"
                        table.cell(row_idx, total_idx).text = f"{r_data['当行合计']:.2f}"
                        row_idx += 1

                    table.cell(row_idx, 0).merge(table.cell(row_idx, total_idx - 1)).text = "合计"
                    table.cell(row_idx, total_idx).text = f"{total_sum:.2f}"

                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in paragraph.runs:
                                    run.font.name = '宋体'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体'); run.font.size = Pt(10)

                    merged_doc.add_paragraph("\n支付银行账户信息如下：")
                    binfo = [
                        ("名称", bank_info.get("名称", "未配置")),
                        ("银行类别", bank_info.get("银行类别", "未配置")),
                        ("开户银行名称", bank_info.get("开户银行名称", "未配置")),
                        ("银行账号", bank_info.get("银行账号", "未配置"))
                    ]
                    for k, v in binfo:
                        p_bank = merged_doc.add_paragraph(f"{k}：{v}" if k != "名称" else f"{k}，{v}")
                        for run in p_bank.runs:
                            run.font.name = '宋体'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                    merged_doc.add_paragraph("\n")
                    sig_name = sys_settings.get("company_signature", "省公众人力部")
                    p_sig = merged_doc.add_paragraph(sig_name)
                    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p_date = merged_doc.add_paragraph(datetime.datetime.now().strftime("%Y年%m月%d日"))
                    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for p in [p_sig, p_date]:
                        for run in p.runs:
                            run.font.name = '宋体'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体'); run.font.size = Pt(12)

            # 将 Word 文件流锁死在本地缓存中
            doc_io = io.BytesIO()
            merged_doc.save(doc_io)
            st.session_state['ss_word_data'] = doc_io.getvalue()
            st.session_state['ss_word_filename'] = f"全量地市结算函_合并打印版_{s_m}至{e_m}.docx"
            st.success("✅ 公对公结算函排版完毕，请点击下方按钮安全下载！")

    # 独立暴露的外层下载按钮
    if 'ss_word_data' in st.session_state:
        st.download_button(
            label=f"📥 2. 点击下载【合并打印版】结算函",
            data=st.session_state['ss_word_data'],
            file_name=st.session_state['ss_word_filename'],
            type="secondary"
        )

    st.divider()
    st.subheader("📚 结算批次台账")
    batch_df = get_settlement_batches_dataframe()
    if batch_df.empty:
        st.caption("生成结算函后，系统会在这里登记唯一批次，避免同一期间重复结算。")
    else:
        batch_display = batch_df.copy()
        batch_display['status'] = batch_display['status'].map(
            SETTLEMENT_BATCH_STATUS_LABELS
        ).fillna(batch_display['status'])
        st.dataframe(
            batch_display[[
                'batch_id', 'period_start', 'period_end', 'branch_name', 'payee_name',
                'total_amount', 'settled_amount', 'status', 'generated_at', 'settled_at', 'voucher_no'
            ]].rename(columns={
                'batch_id': '批次号', 'period_start': '开始月', 'period_end': '结束月',
                'branch_name': '付款地市', 'payee_name': '收款主体',
                'total_amount': '应结金额', 'settled_amount': '已结金额',
                'status': '状态', 'generated_at': '生成时间', 'settled_at': '到账时间',
                'voucher_no': '凭证号'
            }),
            use_container_width=True,
            hide_index=True,
        )
        with st.expander("更新结算到账状态", expanded=False):
            batch_options = batch_df['batch_id'].tolist()
            batch_labels = {
                row['batch_id']: f"{row['period_start']}至{row['period_end']} | {row['branch_name']} → {row['payee_name']}"
                for _, row in batch_df.iterrows()
            }
            with st.form("settlement_status_form"):
                selected_batch = st.selectbox(
                    "结算批次", batch_options, format_func=lambda value: batch_labels[value]
                )
                selected_source = batch_df[batch_df['batch_id'] == selected_batch].iloc[0]
                batch_status_options = ['generated', 'sent', 'confirmed', 'paid']
                batch_status = st.selectbox(
                    "状态", batch_status_options,
                    format_func=lambda value: SETTLEMENT_BATCH_STATUS_LABELS[value],
                )
                settled_amount = st.number_input(
                    "已结金额", min_value=0.0,
                    value=float(selected_source['settled_amount'] or 0.0),
                    step=0.01
                )
                voucher_no = st.text_input("到账/会计凭证号", value=str(selected_source['voucher_no'] or ''))
                if st.form_submit_button("保存批次状态", type="primary"):
                    ok, msg = update_settlement_batch_status(
                        selected_batch, batch_status, settled_amount, voucher_no
                    )
                    if ok:
                        st.success(msg)
                    else:
                        st.error(msg)

# ------------------------------------------------------------------------------
# Tab 3: 全局规则与参数配置
# ------------------------------------------------------------------------------
with tab3:
    st.subheader("🛠️ 全局算力引擎参数设置")
    st.info("💡 此处的设置将直接决定全公司数千条社保数据的最终核算金额。")

    c_y, c_e = st.columns(2)
    with c_y: target_year = st.selectbox("📅 规则生效年度", ["2024", "2025", "2026", "2027", "2028"], index=1)
    with c_e: target_entity = st.selectbox("🏢 适用法人主体", ["全量设置", "省公众", "中电数智", "省公司"])

    fetch_entity = "省公众" if "全量" in target_entity else target_entity
    curr = get_policy_rules(target_year, fetch_entity)

    if curr: st.success(f"已加载 {target_year} 年度 【{fetch_entity}】 的历史配置，可直接修改。")
    else: st.info(f"【{fetch_entity}】 在 {target_year} 年度尚无配置记录，请录入参数后保存。")

    if st.session_state.get('show_confirm', False):
        st.warning("⚠️ 警告：您即将修改底层财务参数，请核对变更情况！")
        if st.button("🔥 确认无误，强行写入数据库", type="primary", key="confirm_upsert_btn"):
            params = st.session_state.get('pending_params')
            if params:
                success, msg = upsert_policy_rules(params, is_all_entities=("全量" in params[1]))
                if success:
                    st.success(msg)
                    st.session_state['show_confirm'] = False
                    st.rerun()
                else: st.error(f"🚨 数据库写入被底层拦截！真实原因: {msg}")
        if st.button("❌ 撤销修改，返回编辑", key="cancel_upsert_btn"):
            st.session_state['show_confirm'] = False
            st.rerun()
        st.divider()

    with st.form("policy_rules_form"):
        st.write(f"**【{target_year} 年度】算法控制总开关**")
        c_mode, c_fund, c_med = st.columns(3)
        with c_mode:
            # [核心修复3] 补充中文翻译映射字典，彻底告别底层代码暴露
            r_keys = ['exact', 'round_to_yuan', 'round_to_ten', 'floor_to_ten']
            r_map = {
                'exact': '精确到分 (不取整)',
                'round_to_yuan': '四舍五入到元',
                'round_to_ten': '四舍五入到十元',
                'floor_to_ten': '向下取整到十元 (见角进元等)'
            }
            cur_round = curr.get('rounding_mode', 'round_to_yuan')
            # 引入 format_func 让界面显示中文，底层仍传英文
            sel_round = st.selectbox("社保取整规则", options=r_keys, format_func=lambda x: r_map[x], index=r_keys.index(cur_round) if cur_round in r_keys else 1)

        with c_fund:
            # [核心修复4] 公积金算法的中文翻译字典
            f_keys = ['independent', 'reverse_from_ss']
            f_map = {
                'independent': '独立计算 (基数×比例)',
                'reverse_from_ss': '反推法 (企个相加逢元进十等)'
            }
            cur_fund = curr.get('fund_calc_method', 'reverse_from_ss')
            sel_fund = st.selectbox("公积金特殊算法", options=f_keys, format_func=lambda x: f_map[x], index=f_keys.index(cur_fund) if cur_fund in f_keys else 1)

        with c_med:
            med_serious = st.number_input("大病医疗个人固定扣款", value=safe_float(curr.get('medical_serious_fix', 7.0)), step=1.0)

        st.divider()
        def render_ins_row(label, prefix, has_pers=True, has_limit=True):
            cols = st.columns([1.5, 2, 2, 2, 2])
            cols[0].markdown(f"**{label}**")
            up = cols[1].number_input(f"{label}封顶", value=safe_float(curr.get(f'{prefix}_upper')), step=100.0, label_visibility="collapsed") if has_limit else 0.0
            lw = cols[2].number_input(f"{label}保底", value=safe_float(curr.get(f'{prefix}_lower')), step=100.0, label_visibility="collapsed") if has_limit else 0.0
            cr = cols[3].number_input(f"{label}企%", value=safe_float(curr.get(f'{prefix}_comp_rate')), step=0.01, format="%.4f", label_visibility="collapsed")
            pr = cols[4].number_input(f"{label}个%", value=safe_float(curr.get(f'{prefix}_pers_rate')), step=0.01, format="%.4f", label_visibility="collapsed") if has_pers else 0.0
            return up, lw, cr, pr

        p_up, p_lw, p_cr, p_pr = render_ins_row("养老保险", "pension")
        m_up, m_lw, m_cr, m_pr = render_ins_row("医疗保险", "medical")
        u_up, u_lw, u_cr, u_pr = render_ins_row("失业保险", "unemp")
        i_up, i_lw, i_cr, _ = render_ins_row("工伤保险", "injury", has_pers=False)
        mat_up, mat_lw, mat_cr, _ = render_ins_row("生育保险", "maternity", has_pers=False)
        f_up, f_lw, f_cr, f_pr = render_ins_row("住房公积金(官方线)", "fund")

        cols_soe = st.columns([1.5, 2, 2, 2, 2])
        cols_soe[0].markdown("**↳ 内部实际执行线**")
        f_soe_up = cols_soe[1].number_input("内部封顶", value=safe_float(curr.get('fund_soe_upper')), step=100.0, label_visibility="collapsed")
        f_soe_lw = cols_soe[2].number_input("内部保底", value=safe_float(curr.get('fund_soe_lower')), step=100.0, label_visibility="collapsed")

        _, _, a_cr, a_pr = render_ins_row("企业年金", "annuity", has_limit=False)

        if st.form_submit_button("🔍 对比并预览修改", type="primary"):
            st.session_state['pending_params'] = (
                target_year, target_entity, sel_round, sel_fund, med_serious,
                p_up, p_lw, p_cr, p_pr, m_up, m_lw, m_cr, m_pr,
                u_up, u_lw, u_cr, u_pr, i_up, i_lw, i_cr,
                mat_up, mat_lw, mat_cr, f_up, f_lw, f_cr, f_pr,
                a_cr, a_pr, f_soe_up, f_soe_lw
            )
            st.session_state['show_confirm'] = True
            st.rerun()

# ------------------------------------------------------------------------------
# Tab 4: 历史账单导入 (冷启动专区)
# ------------------------------------------------------------------------------
with tab4:
    st.subheader("📥 历史月度死账导入引擎 (冷启动与时光机)")
    st.warning("⚠️ 极度危险操作！此功能用于系统上线前的历史账单补录。利用“时光机”功能，可自动抓取最近月份的数据作为底稿，免去手动查工号的痛苦。")

    conn = _get_db_connection()
    try:
        archived_months_df = pd.read_sql_query("SELECT DISTINCT cost_month FROM ss_monthly_records ORDER BY cost_month DESC", conn)
        archived_months = archived_months_df['cost_month'].tolist()
    except Exception:
        archived_months = []
    finally:
        conn.close()

    hc1, hc2 = st.columns(2)
    with hc1:
        target_hist_month = st.text_input("📅 1. 设定你要补录的目标月份 (如: 2026-02)", value="2026-02")

        ref_month = None
        if archived_months:
            try:
                def ym_to_int(ym_str):
                    y, m = map(int, str(ym_str).split('-'))
                    return y * 12 + m
                target_val = ym_to_int(target_hist_month)
                ref_month = min(archived_months, key=lambda x: abs(ym_to_int(x) - target_val))
            except Exception:
                ref_month = archived_months[0]

        # [修复点：防跳跃] 历史底稿的生成与下载分离
        if st.button("🚀 2. 提取历史框架 (生成预填底稿)", type="primary"):
            conn = _get_db_connection()
            if ref_month:
                st.info(f"💡 探针已激活：系统侦测到最近的关联账单为 {ref_month}，正在为您完美克隆该月数据作为底稿...")
                clone_query = """
                    SELECT r.*, e.name AS '姓名'
                    FROM ss_monthly_records r
                    LEFT JOIN employees e ON r.emp_id = e.emp_id
                    WHERE r.cost_month = ?
                """
                ref_df = pd.read_sql_query(clone_query, conn, params=[ref_month])
            else:
                st.info("💡 探针已激活：系统暂无历史账单，正在从现有人事档案中提取名单生成空白底稿...")
                clone_query = """
                    SELECT e.emp_id, e.name AS '姓名', IFNULL(m.cost_center, '本级') AS cost_center
                    FROM employees e
                    LEFT JOIN ss_emp_matrix m ON e.emp_id = m.emp_id
                    WHERE e.status IN ('在职', '挂靠人员')
                """
                ref_df = pd.read_sql_query(clone_query, conn)
            conn.close()

            hist_cols = [
                '核算月份(YYYY-MM)', '工号', '姓名', '财务归属(成本中心)',
                '养老_企金额', '养老_个金额', '养老_通道(如:省公众)',
                '医疗_企金额', '医疗_个金额', '大病_个固定', '医疗_通道(如:省公司)',
                '失业_企金额', '失业_个金额', '失业_通道(如:省公众)',
                '工伤_企金额', '工伤_通道(如:省公众)',
                '生育_企金额', '生育_通道(如:省公司)',
                '公积金_企金额', '公积金_个金额', '公积金_通道(如:省公众)',
                '年金_企金额', '年金_个金额', '年金_通道(如:省公司)'
            ]

            output_data = []
            for _, row in ref_df.iterrows():
                out_row = {col: 0.0 for col in hist_cols}
                out_row['核算月份(YYYY-MM)'] = target_hist_month
                out_row['工号'] = row.get('emp_id', '')
                out_row['姓名'] = row.get('姓名', '')
                out_row['财务归属(成本中心)'] = row.get('cost_center', '本级')

                if ref_month:
                    out_row['养老_企金额'] = row.get('pension_comp', 0.0)
                    out_row['养老_个金额'] = row.get('pension_pers', 0.0)
                    out_row['养老_通道(如:省公众)'] = row.get('pension_route', '')
                    out_row['医疗_企金额'] = row.get('medical_comp', 0.0)
                    out_row['医疗_个金额'] = row.get('medical_pers', 0.0)
                    out_row['大病_个固定'] = row.get('medical_serious_pers', 0.0)
                    out_row['医疗_通道(如:省公司)'] = row.get('medical_route', '')
                    out_row['失业_企金额'] = row.get('unemp_comp', 0.0)
                    out_row['失业_个金额'] = row.get('unemp_pers', 0.0)
                    out_row['失业_通道(如:省公众)'] = row.get('unemp_route', '')
                    out_row['工伤_企金额'] = row.get('injury_comp', 0.0)
                    out_row['工伤_通道(如:省公众)'] = row.get('injury_route', '')
                    out_row['生育_企金额'] = row.get('maternity_comp', 0.0)
                    out_row['生育_通道(如:省公司)'] = row.get('maternity_route', '')
                    out_row['公积金_企金额'] = row.get('fund_comp', 0.0)
                    out_row['公积金_个金额'] = row.get('fund_pers', 0.0)
                    out_row['公积金_通道(如:省公众)'] = row.get('fund_route', '')
                    out_row['年金_企金额'] = row.get('annuity_comp', 0.0)
                    out_row['年金_个金额'] = row.get('annuity_pers', 0.0)
                    out_row['年金_通道(如:省公司)'] = row.get('annuity_route', '')

                output_data.append(out_row)

            hist_template = pd.DataFrame(output_data)
            hist_buffer = io.BytesIO()
            with pd.ExcelWriter(hist_buffer, engine='xlsxwriter') as writer:
                hist_template.to_excel(writer, index=False)
                writer.sheets['Sheet1'].set_column('A:X', 18)

            st.session_state['ss_hist_data'] = hist_buffer.getvalue()
            st.session_state['ss_hist_filename'] = f"{target_hist_month}_历史社保智能补录底稿.xlsx"
            st.success("✅ 底稿结构已提取完毕，请点击下方按钮安全下载！")

        # 悬挂在外部的隔离下载按钮
        if 'ss_hist_data' in st.session_state:
            st.download_button(
                label="📥 3. 点击下载已预填好的历史账单底稿",
                data=st.session_state['ss_hist_data'],
                file_name=st.session_state['ss_hist_filename'],
                type="secondary"
            )

    with hc2:
        hist_file = st.file_uploader("📤 4. 上传核对完毕的历史账单明细", type=["xlsx", "csv"])
        if hist_file and st.button("🚨 5. 强行覆写历史月度账单入库", type="primary"):
            try:
                h_df = pd.read_csv(hist_file) if hist_file.name.endswith('.csv') else pd.read_excel(hist_file)
                conn = _get_db_connection()
                cursor = conn.cursor()

                upsert_sql = """
                    INSERT INTO ss_monthly_records (
                        record_id, cost_month, emp_id, cost_center,
                        pension_comp, pension_pers, pension_route,
                        medical_comp, medical_pers, medical_serious_pers, medical_route,
                        unemp_comp, unemp_pers, unemp_route,
                        injury_comp, injury_route,
                        maternity_comp, maternity_route,
                        fund_comp, fund_pers, fund_route,
                        annuity_comp, annuity_pers, annuity_route
                    ) VALUES (?,?,?,?, ?,?,?, ?,?,?,?, ?,?,?, ?,?, ?,?, ?,?,?, ?,?,?)
                    ON CONFLICT(cost_month, emp_id) DO UPDATE SET
                        cost_center=excluded.cost_center,
                        pension_comp=excluded.pension_comp, pension_pers=excluded.pension_pers, pension_route=excluded.pension_route,
                        medical_comp=excluded.medical_comp, medical_pers=excluded.medical_pers, medical_serious_pers=excluded.medical_serious_pers, medical_route=excluded.medical_route,
                        unemp_comp=excluded.unemp_comp, unemp_pers=excluded.unemp_pers, unemp_route=excluded.unemp_route,
                        injury_comp=excluded.injury_comp, injury_route=excluded.injury_route,
                        maternity_comp=excluded.maternity_comp, maternity_route=excluded.maternity_route,
                        fund_comp=excluded.fund_comp, fund_pers=excluded.fund_pers, fund_route=excluded.fund_route,
                        annuity_comp=excluded.annuity_comp, annuity_pers=excluded.annuity_pers, annuity_route=excluded.annuity_route
                """
                count = 0
                for _, row in h_df.iterrows():
                    h_month = str(row.get('核算月份(YYYY-MM)', '')).strip()
                    eid = str(row.get('工号', '')).replace('.0', '').strip()
                    if not h_month or not eid or h_month == 'nan' or eid == 'nan': continue

                    rec_id = f"{h_month}_{eid}"
                    cc = str(row.get('财务归属(成本中心)', '本级')).strip()

                    cursor.execute(upsert_sql, (
                        rec_id, h_month, eid, cc,
                        safe_float(row.get('养老_企金额')), safe_float(row.get('养老_个金额')), str(row.get('养老_通道(如:省公众)', 'None')),
                        safe_float(row.get('医疗_企金额')), safe_float(row.get('医疗_个金额')), safe_float(row.get('大病_个固定')), str(row.get('医疗_通道(如:省公司)', 'None')),
                        safe_float(row.get('失业_企金额')), safe_float(row.get('失业_个金额')), str(row.get('失业_通道(如:省公众)', 'None')),
                        safe_float(row.get('工伤_企金额')), str(row.get('工伤_通道(如:省公众)', 'None')),
                        safe_float(row.get('生育_企金额')), str(row.get('生育_通道(如:省公司)', 'None')),
                        safe_float(row.get('公积金_企金额')), safe_float(row.get('公积金_个金额')), str(row.get('公积金_通道(如:省公众)', 'None')),
                        safe_float(row.get('年金_企金额')), safe_float(row.get('年金_个金额')), str(row.get('年金_通道(如:省公司)', 'None'))
                    ))
                    count += 1
                conn.commit()
                snapshot_counts = backfill_relationship_snapshots()
                st.success(
                    f"✅ 历史账单成功覆盖/写入 {count} 条；"
                    f"关系与险种快照已同步 {snapshot_counts['social_items']} 条。"
                )
            except Exception as e:
                st.error(f"❌ 导入崩溃: {e}")
            finally:
                if 'conn' in locals(): conn.close()

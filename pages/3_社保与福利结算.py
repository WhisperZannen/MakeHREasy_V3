# ==============================================================================
# 文件路径: pages/3_社保与福利结算.py
# 功能描述: 多主体社保与福利结算中枢 (MVC 架构前端 UI 层 - 财务隔离终极版)
# 核心改造:
#   1. Tab 1: 完美分离 199 与 7。
#   2. Tab 2: 彻底修复 Grouper 1D 崩溃。实装 5 大独立财务提取单引擎（总览与月度拆分）。
#   3. Tab 2 第二部分：[核心重构] 将零散的 ZIP 下载彻底替换为“一键生成全量合并 Word”。
#   4. Tab 4: 历史补录支持时间感知倒推（智能寻轨引擎）。
# ==============================================================================

import streamlit as st
import pandas as pd
import io
import uuid
import zipfile
import datetime
import json
import os

# 引入公文渲染引擎
try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    st.error("🚨 缺少 Word 生成引擎！请在终端运行：pip install python-docx")
    st.stop()

from modules.core_social_security import (
    get_policy_rules,
    upsert_policy_rules,
    _get_db_connection,
    batch_update_emp_matrix
)

st.set_page_config(page_title="社保与福利结算", layout="wide")

# ==============================================================================
# [配置中枢] Settings.json 动态热加载与自愈引擎
# ==============================================================================
def load_settings():
    current_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(current_dir)
    settings_path = os.path.join(project_root, 'settings.json')

    if not os.path.exists(settings_path):
        default_settings = {
            "bank_accounts": {
                "省公众": {
                    "名称": "湖北公众信息产业有限责任公司",
                    "银行类别": "交通银行",
                    "开户银行名称": "交通银行北京市分行营业部",
                    "银行账号": "990204011701001401"
                },
                "中电数智": {
                    "名称": "中电信数智科技有限公司湖北分公司",
                    "银行类别": "交通银行",
                    "开户银行名称": "交通银行北京市分行营业部",
                    "银行账号": "990204011701007001"
                }
            },
            "company_signature": "省公众人力部"
        }
        with open(settings_path, 'w', encoding='utf-8') as f:
            json.dump(default_settings, f, ensure_ascii=False, indent=4)
        return default_settings

    with open(settings_path, 'r', encoding='utf-8') as f:
        return json.load(f)

# ==============================================================================
# [核心防御机制] 空值清洗器
# ==============================================================================
def safe_float(val, default=0.0):
    try:
        if pd.notna(val) and val is not None and str(val).strip() != '':
            return float(val)
        return default
    except Exception:
        return default

if 'show_confirm' not in st.session_state: st.session_state['show_confirm'] = False
if 'pending_params' not in st.session_state: st.session_state['pending_params'] = None

st.title("🛡️ 社保与福利结算中心")
st.caption("核心业务流向：当月基数备料 ➡️ 理论核算与补缴对账 ➡️ 跨主体结算与公对公要款 ➡️ 引擎底座配置")

tab1, tab2, tab3, tab4 = st.tabs(["🧮 当月社保沙盘 (含补缴)", "📤 财务提款与公对公结算函", "⚙️ 全局规则与参数配置", "📥 历史账单导入 (冷启动)"])

# ------------------------------------------------------------------------------
# Tab 1: 当月社保沙盘与对账池
# ------------------------------------------------------------------------------
with tab1:
    st.info("💡 业务铁律：先在【第一步】确保所有人基数就绪，再在【第二步】跑出当期理论账单，最后在【第三步】补录官方滞纳金与历史补缴差额。")

    calc_month = st.text_input("📅 输入当前核算工作月份 (格式: YYYY-MM，如 2026-03)", value="2026-03", max_chars=7)

    st.subheader("🛠️ 第一步：基数初始化与特例抢救")
    conn = _get_db_connection()
    try:
        detect_sql = """
            SELECT 
                e.emp_id AS '工号', e.name AS '姓名', d.dept_name AS '部门', e.status AS '人事状态',
                IFNULL(m.cost_center, '本级') AS '财务归属', IFNULL(m.base_salary_avg, 0.0) AS '已录入原始基数',
                IFNULL(m.fund_base_avg, 0.0) AS '独立公积金基数(选填)',
                IFNULL(m.pension_enabled, 1) AS '养老参保(1是0否)', IFNULL(m.pension_account, '省公众') AS '养老缴纳主体',
                IFNULL(m.medical_enabled, 1) AS '医疗参保(1是0否)', IFNULL(m.medical_account, '省公司') AS '医疗缴纳主体',
                7.0 AS '大病统筹(个人固定)',
                IFNULL(m.unemp_enabled, 1) AS '失业参保(1是0否)', IFNULL(m.unemp_account, '省公众') AS '失业缴纳主体',
                IFNULL(m.injury_enabled, 1) AS '工伤参保(1是0否)', IFNULL(m.injury_account, '省公众') AS '工伤缴纳主体',
                IFNULL(m.maternity_enabled, 1) AS '生育参保(1是0否)', IFNULL(m.maternity_account, '省公司') AS '生育缴纳主体',
                IFNULL(m.fund_enabled, 1) AS '公积金参保(1是0否)', IFNULL(m.fund_account, '省公众') AS '公积金缴纳主体',
                IFNULL(m.annuity_enabled, 0) AS '年金参保(1是0否)', IFNULL(m.annuity_account, '省公司') AS '年金缴纳主体'
            FROM employees e
            LEFT JOIN departments d ON e.dept_id = d.dept_id
            LEFT JOIN ss_emp_matrix m ON e.emp_id = m.emp_id
            WHERE e.status IN ('在职', '挂靠人员')
        """
        roster_df = pd.read_sql_query(detect_sql, conn)
    finally:
        conn.close()

    c_down, c_up = st.columns(2)
    with c_down:
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
            roster_df.to_excel(writer, index=False, sheet_name='基数名单')
        st.download_button("📥 1. 下载当期参保人员基数名单", data=buffer.getvalue(), file_name=f"全员基数表_{calc_month}.xlsx", mime="application/vnd.ms-excel")
    with c_up:
        uploaded_file = st.file_uploader("📤 2. 上传填好基数/修改过开关的 Excel", type=["xlsx", "csv"], label_visibility="collapsed")
        if uploaded_file is not None:
            if st.session_state.get('last_processed_file_id') != uploaded_file.file_id:
                try:
                    upload_df = pd.read_csv(uploaded_file) if uploaded_file.name.endswith('.csv') else pd.read_excel(uploaded_file)
                    success, msg = batch_update_emp_matrix(upload_df)
                    if success:
                        st.session_state['last_processed_file_id'] = uploaded_file.file_id
                        st.rerun()
                    else: st.error(msg)
                except Exception as e: st.error(f"❌ 读取文件失败: {e}")
            else: st.success("✅ 基数已成功灌入引擎，可以开始核算！")

    st.divider()

    st.subheader("🧮 第二步：本月正常参保核算")
    rule_year_to_use = st.selectbox("⚙️ 选择本次套用的【规则年度】(如次年6月前沿用上年规则)", ["2024", "2025", "2026", "2027", "2028"], index=1)

    if st.button("🚀 启动引擎，生成当期理论账单", type="primary"):
        missing_count = len(roster_df[roster_df['已录入原始基数'] == 0.0])
        if missing_count > 0:
            st.error(f"🚨 警告：发现 {missing_count} 名员工基数为 0，请在第一步补充完整！")
        else:
            all_bills = []
            from modules.core_social_security import calculate_complete_bill
            for _, row in roster_df.iterrows():
                all_bills.append(calculate_complete_bill(row.to_dict(), rule_year_to_use))
            st.session_state['temp_bills'] = pd.DataFrame(all_bills)

    if 'temp_bills' in st.session_state:
        raw_df_preview = st.session_state['temp_bills']
        export_df = raw_df_preview.copy()

        cols_to_drop = ['injury_个', 'maternity_个']
        export_df = export_df.drop(columns=[c for c in cols_to_drop if c in export_df.columns])

        audit_rename_map = {
            'pension_企': '养老(企业)', 'pension_个': '养老(个人)', 'pension_route': '养老缴纳主体',
            'medical_企': '医疗(企业)', 'medical_个': '医疗(个人)', 'medical_route': '医疗缴纳主体',
            'medical_serious_个': '大病(个人)', 'medical_serious_pers': '大病(个人)', '大病_个': '大病(个人)',
            'unemp_企': '失业(企业)', 'unemp_个': '失业(个人)', 'unemp_route': '失业缴纳主体',
            'injury_企': '工伤(企业)', 'injury_route': '工伤缴纳主体',
            'maternity_企': '生育(企业)', 'maternity_route': '生育缴纳主体',
            'fund_企': '公积金(企业)', 'fund_个': '公积金(个人)', 'fund_route': '公积金缴纳主体',
            'annuity_企': '年金(企业)', 'annuity_个': '年金(个人)', 'annuity_route': '年金缴纳主体'
        }
        export_df = export_df.rename(columns=audit_rename_map)

        if '大病(个人)' not in export_df.columns:
            export_df['大病(个人)'] = 0.0

        ordered_front_cols = ['工号', '姓名', '财务归属', '合计企业缴纳', '合计个人扣款']
        detail_cols = [c for c in export_df.columns if c not in ordered_front_cols]
        export_df = export_df[ordered_front_cols + detail_cols]

        search_query = st.text_input("🔍 抽查指定员工 (输入姓名或工号进行过滤审核)", "")
        display_df = export_df
        if search_query:
            display_df = display_df[
                display_df['姓名'].str.contains(search_query, na=False) |
                display_df['工号'].str.contains(search_query, na=False)
            ]

        st.dataframe(display_df[['工号', '姓名', '财务归属', '合计企业缴纳', '合计个人扣款']], use_container_width=True, hide_index=True)

        st.write("---")
        c_audit, c_save = st.columns(2)

        with c_audit:
            buffer_audit = io.BytesIO()
            with pd.ExcelWriter(buffer_audit, engine='xlsxwriter') as writer:
                export_df.to_excel(writer, index=False, sheet_name='当月核算全量底稿')

            st.download_button(
                label="📥 1. 下载全量明细底稿 (全中文，供线下复核)",
                data=buffer_audit.getvalue(),
                file_name=f"当期核算明细底稿_{calc_month}.xlsx",
                type="secondary"
            )

        with c_save:
            if st.button("💾 2. 线下复核无误，将当期明细固化入库", type="primary"):
                from modules.core_social_security import save_monthly_ss_records
                success, msg = save_monthly_ss_records(raw_df_preview, calc_month)
                if success:
                    st.success(msg)
                    del st.session_state['temp_bills']
                    st.rerun()
                else:
                    st.error(msg)

    st.subheader("📥 第三步：补缴与滞纳金手工入账 (对齐官方核定单)")
    st.write("🔧 遇到历史跨月补缴、滞纳金等系统无法自动推演的账目，请在此按社保局单据直接填报写入。")

    rc1, rc2 = st.columns(2)
    with rc1:
        retro_cols = ['处理月份(即本月)', '工号', '补缴起始月', '补缴结束月', '补缴险种(必选下拉框)', '企业本金合计', '个人本金合计', '企业承担滞纳金', '备注']
        retro_template = pd.DataFrame(columns=retro_cols)
        retro_buffer = io.BytesIO()
        with pd.ExcelWriter(retro_buffer, engine='xlsxwriter') as writer:
            retro_template.to_excel(writer, index=False, sheet_name='补缴模板')
            workbook  = writer.book
            worksheet = writer.sheets['补缴模板']
            worksheet.data_validation('E2:E1048576', {
                'validate': 'list',
                'source': ['养老保险', '医疗保险', '大病医疗', '失业保险', '工伤保险', '生育保险', '住房公积金', '企业年金']
            })
            header_format = workbook.add_format({'bg_color': '#FFC7CE', 'font_color': '#9C0006', 'bold': True})
            worksheet.write('E1', '补缴险种(必选下拉框)', header_format)
            worksheet.set_column('A:I', 18)

        st.download_button("📥 下载补缴与滞纳金导入模板 (含严格限制)", data=retro_buffer.getvalue(), file_name=f"补缴导入模板_{calc_month}.xlsx")

    with rc2:
        retro_file = st.file_uploader("📤 上传已填好的补缴核定单", type=["xlsx", "csv"], label_visibility="collapsed")
        if retro_file and st.button("💾 将补缴数据强行入库"):
            try:
                r_df = pd.read_csv(retro_file) if retro_file.name.endswith('.csv') else pd.read_excel(retro_file)
                conn = _get_db_connection()
                cursor = conn.cursor()
                sql = """
                    INSERT INTO ss_retroactive_records (
                        retro_id, process_month, emp_id, target_start_month, target_end_month, retro_type,
                        total_comp_retro, total_pers_retro, late_fee, remarks
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """
                count = 0
                for _, row in r_df.iterrows():
                    eid = str(row.get('工号', '')).replace('.0', '').strip()
                    if not eid or eid == 'nan': continue
                    cursor.execute(sql, (
                        str(uuid.uuid4())[:12], str(row.get('处理月份(即本月)', calc_month)).strip(), eid,
                        str(row.get('补缴起始月', '')), str(row.get('补缴结束月', '')),
                        str(row.get('补缴险种(必选下拉框)', '未知险种')),
                        safe_float(row.get('企业本金合计', 0.0)), safe_float(row.get('个人本金合计', 0.0)),
                        safe_float(row.get('企业承担滞纳金', 0.0)), str(row.get('备注', ''))
                    ))
                    count += 1
                conn.commit()
                st.success(f"✅ 成功将 {count} 笔特殊补缴与滞纳金记录封印入库！后续台账与工资计算将自动识别叠加。")
            except Exception as e:
                st.error(f"❌ 写入补缴表失败: {e}")
            finally:
                if 'conn' in locals(): conn.close()

# ------------------------------------------------------------------------------
# Tab 2: 财务输出中心 (对内账单与公对公 Word 结算函)
# ------------------------------------------------------------------------------
with tab2:
    st.subheader("📤 第一部分：对内审批提款单 (跨期多表智能打包)")
    st.info("💡 财务走账专用。严格按【缴费主体+险种】物理隔离出 5 个独立的 Excel。多月提取时自动生成【总览】与月度明细 Sheet。")

    conn = _get_db_connection()
    try:
        available_months = pd.read_sql_query("SELECT DISTINCT cost_month FROM ss_monthly_records ORDER BY cost_month DESC", conn)['cost_month'].tolist()
    except Exception:
        available_months = []
    finally:
        conn.close()

    ic1, ic2 = st.columns(2)
    with ic1:
        int_start_month = st.selectbox("📅 对内请款起始月", options=available_months if available_months else ["无数据"], key='int_start', index=len(available_months)-1 if available_months else 0)
    with ic2:
        int_end_month = st.selectbox("📅 对内请款结束月", options=available_months if available_months else ["无数据"], key='int_end', index=0 if available_months else 0)

    if st.button("🚀 极速生成并打包对内提款单 (ZIP)", type="primary") and int_start_month != "无数据":
        s_m, e_m = min(int_start_month, int_end_month), max(int_start_month, int_end_month)
        selected_months = [m for m in available_months if s_m <= m <= e_m]
        selected_months.sort()

        conn = _get_db_connection()

        query = """
            SELECT r.*, e.name AS '姓名'
            FROM ss_monthly_records r 
            LEFT JOIN employees e ON r.emp_id = e.emp_id 
            WHERE r.cost_month >= ? AND r.cost_month <= ?
        """
        raw_df = pd.read_sql_query(query, conn, params=[s_m, e_m])

        retro_query = """
            SELECT r.*, e.name AS '姓名', IFNULL(m.cost_center, '本级') AS 'cost_center'
            FROM ss_retroactive_records r
            LEFT JOIN employees e ON r.emp_id = e.emp_id
            LEFT JOIN ss_emp_matrix m ON r.emp_id = m.emp_id
            WHERE r.process_month >= ? AND r.process_month <= ?
        """
        retro_df = pd.read_sql_query(retro_query, conn, params=[s_m, e_m])
        conn.close()

        if not raw_df.empty or not retro_df.empty:
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:

                rename_map = {
                    'emp_id': '工号', 'cost_center': '财务归属',
                    'pension_comp': '养老(企业)', 'pension_pers': '养老(个人)',
                    'medical_comp': '医疗(企业)', 'medical_pers': '医疗(个人)', 'medical_serious_pers': '大病(个人)',
                    'unemp_comp': '失业(企业)', 'unemp_pers': '失业(个人)',
                    'injury_comp': '工伤(企业)', 'maternity_comp': '生育(企业)',
                    'fund_comp': '公积金(企业)', 'fund_pers': '公积金(个人)',
                    'annuity_comp': '年金(企业)', 'annuity_pers': '年金(个人)'
                }

                channel_configs = [
                    {'name': '1.中电数智(五险两金综合)', 'route': '中电数智', 'items': ['pension', 'medical', 'unemp', 'injury', 'maternity', 'fund', 'annuity']},
                    {'name': '2.省公司(年金)', 'route': '省公司', 'items': ['annuity']},
                    {'name': '3.省公司(医疗_生育_工伤)', 'route': '省公司', 'items': ['medical', 'maternity', 'injury']},
                    {'name': '4.省公众(公积金)', 'route': '省公众', 'items': ['fund']},
                    {'name': '5.省公众(养老_失业_工伤)', 'route': '省公众', 'items': ['pension', 'unemp', 'injury']}
                ]

                all_insurance_items = ['pension', 'medical', 'unemp', 'injury', 'maternity', 'fund', 'annuity']

                if not raw_df.empty:
                    for config in channel_configs:
                        df_channel = raw_df.copy()

                        for it in all_insurance_items:
                            if it not in config['items']:
                                if f'{it}_comp' in df_channel.columns: df_channel[f'{it}_comp'] = 0.0
                                if f'{it}_pers' in df_channel.columns: df_channel[f'{it}_pers'] = 0.0
                                if it == 'medical' and 'medical_serious_pers' in df_channel.columns: df_channel['medical_serious_pers'] = 0.0
                            else:
                                mask = df_channel[f'{it}_route'] == config['route']
                                if f'{it}_comp' in df_channel.columns: df_channel.loc[~mask, f'{it}_comp'] = 0.0
                                if f'{it}_pers' in df_channel.columns: df_channel.loc[~mask, f'{it}_pers'] = 0.0
                                if it == 'medical' and 'medical_serious_pers' in df_channel.columns: df_channel.loc[~mask, 'medical_serious_pers'] = 0.0

                        money_cols = [c for c in df_channel.columns if c.endswith('_comp') or c.endswith('_pers')]
                        df_channel['__row_sum__'] = df_channel[money_cols].sum(axis=1)
                        df_channel = df_channel[df_channel['__row_sum__'] > 0].drop(columns=['__row_sum__'])

                        if not df_channel.empty:
                            excel_io = io.BytesIO()
                            with pd.ExcelWriter(excel_io, engine='xlsxwriter') as writer:
                                group_cols = ['emp_id', '姓名', 'cost_center']
                                active_sum_cols = [c for c in money_cols if df_channel[c].sum() > 0]

                                if active_sum_cols:
                                    df_sum = df_channel.groupby(group_cols)[active_sum_cols].sum().reset_index()
                                    export_cols = group_cols + active_sum_cols
                                    df_export = df_sum[export_cols].rename(columns=rename_map)
                                    df_export.to_excel(writer, index=False, sheet_name="总览")

                                for month in selected_months:
                                    df_month = df_channel[df_channel['cost_month'] == month]
                                    if not df_month.empty:
                                        m_active_cols = [c for c in money_cols if df_month[c].sum() > 0]
                                        if m_active_cols:
                                            export_cols = group_cols + m_active_cols
                                            df_export = df_month[export_cols].rename(columns=rename_map)
                                            df_export.to_excel(writer, index=False, sheet_name=month)

                            zf.writestr(f"{config['name']}_{s_m}至{e_m}.xlsx", excel_io.getvalue())

                if not retro_df.empty:
                    retro_map = {
                        'process_month': '处理月份', 'emp_id': '工号', 'retro_type': '补缴险种',
                        'target_start_month': '补缴起始月', 'target_end_month': '补缴结束月',
                        'total_comp_retro': '企业本金', 'total_pers_retro': '个人本金',
                        'late_fee': '滞纳金(异常支出)', 'remarks': '产生原因(备注)', 'cost_center': '财务归属'
                    }
                    retro_cols = ['处理月份', '工号', '姓名', '财务归属', '补缴险种', '补缴起始月', '补缴结束月', '企业本金', '个人本金', '滞纳金(异常支出)', '产生原因(备注)']
                    df_retro_export = retro_df.rename(columns=retro_map)[retro_cols]

                    excel_io = io.BytesIO()
                    with pd.ExcelWriter(excel_io, engine='xlsxwriter') as writer:
                        df_retro_export.to_excel(writer, index=False, sheet_name="异常款项专项审批")
                        workbook = writer.book
                        worksheet = writer.sheets['异常款项专项审批']
                        alert_format = workbook.add_format({'bg_color': '#FFC7CE', 'font_color': '#9C0006', 'bold': True})
                        worksheet.set_column('J:J', 16, alert_format)
                        worksheet.set_column('K:K', 35)

                    zf.writestr(f"6.异常款项专项审批_{s_m}至{e_m}.xlsx", excel_io.getvalue())

            st.download_button(f"📥 点击下载对内提款单大包 (ZIP)", data=zip_buffer.getvalue(), file_name=f"对内审计提款单合集_{s_m}至{e_m}.zip", type="primary")

    st.divider()

    # ==========================================
    # 第二部分：对外公对公结算函生成引擎 (一键合并打印版)
    # ==========================================
    st.subheader("📜 第二部分：跨期公对公结算函 (Word 动态打包引擎)")
    st.write("🔧 动态侦测全险种特例，自适应伸缩表格。彻底融合【跨月补缴与滞纳金】，精准追溯每一分财务死账。")

    ec1, ec2 = st.columns(2)
    with ec1: start_month = st.selectbox("⏳ 结算起始月", options=available_months if available_months else ["无数据"], key='s_month', index=len(available_months)-1 if available_months else 0)
    with ec2: end_month = st.selectbox("⏳ 结算结束月", options=available_months if available_months else ["无数据"], key='e_month', index=0 if available_months else 0)

    conn = _get_db_connection()
    s_m, e_m = min(start_month, end_month), max(start_month, end_month)
    branch_query = "SELECT DISTINCT cost_center FROM ss_monthly_records WHERE cost_month >= ? AND cost_month <= ? AND cost_center != '本级'"
    avail_branches_df = pd.read_sql_query(branch_query, conn, params=[s_m, e_m])
    conn.close()
    avail_branches = avail_branches_df['cost_center'].dropna().tolist()

    selected_branches = st.multiselect("🏢 勾选需要生成结算函的地市分公司 (默认全选)", options=avail_branches, default=avail_branches)

    # [核心体验升级] 输出单个全量 Word，完全摒弃 ZIP，所有公对公结算单排在一个文件里，点击一次全部打印
    if st.button("🚀 一键生成【合并版】地市结算函 (全量 Word)", type="primary") and start_month != "无数据" and selected_branches:
        sys_settings = load_settings()
        conn = _get_db_connection()
        placeholders = ",".join(["?"] * len(selected_branches))
        ext_params = [s_m, e_m] + selected_branches

        ext_query = f"""
            SELECT r.*, e.name AS '姓名', 
                   COALESCE(NULLIF(m.ss_base_actual, 0.0), NULLIF(m.base_salary_avg, 0.0), 0.0) AS '基数'
            FROM ss_monthly_records r 
            LEFT JOIN employees e ON r.emp_id = e.emp_id 
            LEFT JOIN ss_emp_matrix m ON r.emp_id = m.emp_id
            WHERE r.cost_month >= ? AND r.cost_month <= ? AND r.cost_center IN ({placeholders})
        """
        ext_df = pd.read_sql_query(ext_query, conn, params=ext_params)

        retro_query = f"""
            SELECT r.*, e.name AS '姓名', 
                   IFNULL(m.cost_center, '本级') AS 'cost_center',
                   COALESCE(NULLIF(m.ss_base_actual, 0.0), NULLIF(m.base_salary_avg, 0.0), 0.0) AS '基数',
                   m.pension_account AS 'pension_route', m.medical_account AS 'medical_route',
                   m.unemp_account AS 'unemp_route', m.injury_account AS 'injury_route',
                   m.maternity_account AS 'maternity_route', m.fund_account AS 'fund_route',
                   m.annuity_account AS 'annuity_route'
            FROM ss_retroactive_records r
            LEFT JOIN employees e ON r.emp_id = e.emp_id
            LEFT JOIN ss_emp_matrix m ON r.emp_id = m.emp_id
            WHERE r.process_month >= ? AND r.process_month <= ? AND IFNULL(m.cost_center, '本级') IN ({placeholders})
        """
        retro_df = pd.read_sql_query(retro_query, conn, params=ext_params)
        conn.close()

        retro_map = {'养老保险':'pension', '医疗保险':'medical', '大病医疗':'medical_serious', '失业保险':'unemp', '工伤保险':'injury', '生育保险':'maternity', '住房公积金':'fund', '企业年金':'annuity'}
        normalized_retro = []
        for _, row in retro_df.iterrows():
            prefix = retro_map.get(row.get('retro_type', ''))
            if not prefix: continue

            new_row = {
                'cost_month': f"补缴({row['process_month']})",
                'emp_id': row['emp_id'], '姓名': row['姓名'], 'cost_center': row['cost_center'], '基数': row['基数'],
                f'{prefix}_comp': row['total_comp_retro'], f'{prefix}_pers': row['total_pers_retro'],
                f'{prefix}_route': row.get(f'{prefix}_route', '未知'), 'late_fee': row['late_fee']
            }
            normalized_retro.append(new_row)

        retro_normalized_df = pd.DataFrame(normalized_retro)
        if not retro_normalized_df.empty:
            combined_df = pd.concat([ext_df, retro_normalized_df], ignore_index=True).fillna(0.0)
        else:
            combined_df = ext_df.copy()
            combined_df['late_fee'] = 0.0

        if combined_df.empty:
            st.warning("该时间段内所选分公司没有产生任何费用记录。")
        else:
            merged_doc = Document()
            first_letter = True

            for cc, group in combined_df.groupby('cost_center'):
                routes_used = set()
                for r_col in ['pension_route', 'medical_route', 'unemp_route', 'injury_route', 'maternity_route', 'fund_route', 'annuity_route']:
                    if r_col in group.columns: routes_used.update(group[r_col].astype(str).dropna().unique())
                routes_used.discard(''); routes_used.discard('None'); routes_used.discard('0.0')

                for route_name in routes_used:
                    df_cc_route = group.copy()
                    has_amt = pd.Series([False] * len(df_cc_route))

                    for it in ['pension', 'medical', 'unemp', 'injury', 'maternity', 'fund', 'annuity']:
                        r_col = f'{it}_route'
                        if r_col in df_cc_route.columns:
                            mask = df_cc_route[r_col].astype(str) == route_name
                            df_cc_route.loc[~mask, f'{it}_comp'] = 0.0
                            df_cc_route.loc[~mask, f'{it}_pers'] = 0.0
                            if it == 'medical' and 'medical_serious_pers' in df_cc_route.columns:
                                df_cc_route.loc[~mask, 'medical_serious_pers'] = 0.0

                            c_col_val = df_cc_route[f'{it}_comp'] if f'{it}_comp' in df_cc_route.columns else 0.0
                            p_col_val = df_cc_route[f'{it}_pers'] if f'{it}_pers' in df_cc_route.columns else 0.0
                            has_amt = has_amt | (c_col_val > 0) | (p_col_val > 0)

                    if 'late_fee' in df_cc_route.columns:
                        df_cc_route.loc[~has_amt, 'late_fee'] = 0.0
                        has_amt = has_amt | (df_cc_route['late_fee'] > 0)

                    df_cc_route = df_cc_route[has_amt]
                    if df_cc_route.empty: continue

                    active_cols = []
                    ins_names = {'pension':'养老', 'medical':'医疗', 'medical_serious':'大病', 'unemp':'失业', 'injury':'工伤', 'maternity':'生育', 'fund':'公积金', 'annuity':'年金'}

                    for it in ['pension', 'medical', 'unemp', 'injury', 'maternity', 'fund', 'annuity']:
                        c_sum = df_cc_route[f'{it}_comp'].sum() if f'{it}_comp' in df_cc_route.columns else 0
                        p_sum = df_cc_route[f'{it}_pers'].sum() if f'{it}_pers' in df_cc_route.columns else 0
                        if c_sum > 0 and p_sum > 0: active_cols.append({'id':it, 'name':ins_names[it], 'has_c':True, 'has_p':True})
                        elif c_sum > 0: active_cols.append({'id':it, 'name':ins_names[it], 'has_c':True, 'has_p':False})
                        elif p_sum > 0: active_cols.append({'id':it, 'name':ins_names[it], 'has_c':False, 'has_p':True})

                        if it == 'medical' and 'medical_serious_pers' in df_cc_route.columns:
                            if df_cc_route['medical_serious_pers'].sum() > 0:
                                active_cols.append({'id':'medical_serious', 'name':'大病', 'has_c':False, 'has_p':True})

                    has_late_fee = df_cc_route['late_fee'].sum() > 0 if 'late_fee' in df_cc_route.columns else False

                    row_totals = []
                    for _, r_data in df_cc_route.iterrows():
                        r_tot = 0.0
                        for ac in active_cols:
                            if ac['has_c']: r_tot += r_data.get(f"{ac['id']}_comp", 0.0)
                            if ac['has_p']: r_tot += r_data.get(f"{ac['id']}_pers", 0.0)
                        if has_late_fee: r_tot += r_data.get('late_fee', 0.0)
                        row_totals.append(r_tot)
                    df_cc_route['当行合计'] = row_totals
                    total_sum = sum(row_totals)

                    # [打印分页逻辑] 非首张结算单，强制插入硬分页符，保证每家公司单占一页
                    if not first_letter: merged_doc.add_page_break()
                    first_letter = False

                    # [全量渲染引擎] 将画表代码全部无缝内嵌
                    p_title = merged_doc.add_paragraph()
                    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r_title = p_title.add_run("关于社保代缴的结算说明函")
                    r_title.font.size = Pt(16); r_title.font.bold = True; r_title.font.name = '黑体'
                    r_title._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

                    p_salut = merged_doc.add_paragraph()
                    r_salut = p_salut.add_run(f"{cc}：")
                    r_salut.font.name = '宋体'; r_salut._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                    emp_names = "、".join(df_cc_route['姓名'].unique())
                    if len(emp_names) > 15: emp_names = emp_names[:15] + "等"

                    bank_info = sys_settings.get("bank_accounts", {}).get(route_name, {})
                    account_name = bank_info.get("名称", route_name)

                    ins_str = "、".join([ac['name'] for ac in active_cols])
                    p_body1 = merged_doc.add_paragraph(f"    因业务开展需要，{cc}{emp_names}社保（{ins_str}）暂由{account_name}代缴，代缴金额据实结算。")
                    p_body2 = merged_doc.add_paragraph(f"    从{s_m[:4]}年{s_m[-2:]}月到{e_m[:4]}年{e_m[-2:]}月，代缴金额为{total_sum:.2f}元，明细如下：\n")
                    for p in [p_body1, p_body2]:
                        for run in p.runs:
                            run.font.name = '宋体'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体'); run.font.size = Pt(12)

                    base_headers = ["月份", "公司", "员工姓名", "基数"]
                    num_cols = len(base_headers)
                    for ac in active_cols:
                        if ac['has_c'] and ac['has_p']: num_cols += 2
                        else: num_cols += 1
                    if has_late_fee: num_cols += 1
                    num_cols += 1

                    table = merged_doc.add_table(rows=2 + len(df_cc_route) + 1, cols=num_cols)
                    table.style = 'Table Grid'

                    for i, h in enumerate(base_headers):
                        table.cell(0, i).merge(table.cell(1, i)).text = h

                    col_idx = len(base_headers)
                    for ac in active_cols:
                        if ac['has_c'] and ac['has_p']:
                            table.cell(0, col_idx).merge(table.cell(0, col_idx+1)).text = ac['name']
                            table.cell(1, col_idx).text = "企业"
                            table.cell(1, col_idx+1).text = "个人"
                            ac['c_idx'] = col_idx; ac['p_idx'] = col_idx + 1
                            col_idx += 2
                        elif ac['has_c']:
                            table.cell(0, col_idx).merge(table.cell(1, col_idx)).text = ac['name'] + "(企)"
                            ac['c_idx'] = col_idx; ac['p_idx'] = -1
                            col_idx += 1
                        elif ac['has_p']:
                            table.cell(0, col_idx).merge(table.cell(1, col_idx)).text = ac['name'] + "(个)"
                            ac['c_idx'] = -1; ac['p_idx'] = col_idx
                            col_idx += 1

                    if has_late_fee:
                        table.cell(0, col_idx).merge(table.cell(1, col_idx)).text = "滞纳金"
                        lf_idx = col_idx
                        col_idx += 1

                    table.cell(0, col_idx).merge(table.cell(1, col_idx)).text = "合计"
                    total_idx = col_idx

                    row_idx = 2
                    for _, r_data in df_cc_route.iterrows():
                        month_str = str(r_data['cost_month']).replace('补缴(', '').replace(')', '')
                        if '-' in month_str: month_str = month_str.split('-')[-1] + "月"

                        table.cell(row_idx, 0).text = month_str
                        table.cell(row_idx, 1).text = cc
                        table.cell(row_idx, 2).text = str(r_data['姓名'])
                        table.cell(row_idx, 3).text = str(r_data['基数'])

                        for ac in active_cols:
                            if ac['has_c']: table.cell(row_idx, ac['c_idx']).text = f"{r_data.get(ac['id']+'_comp', 0.0):.2f}"
                            if ac['has_p']: table.cell(row_idx, ac['p_idx']).text = f"{r_data.get(ac['id']+'_pers', 0.0):.2f}"

                        if has_late_fee: table.cell(row_idx, lf_idx).text = f"{r_data.get('late_fee', 0.0):.2f}"
                        table.cell(row_idx, total_idx).text = f"{r_data['当行合计']:.2f}"
                        row_idx += 1

                    table.cell(row_idx, 0).merge(table.cell(row_idx, total_idx - 1)).text = "合计"
                    table.cell(row_idx, total_idx).text = f"{total_sum:.2f}"

                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in paragraph.runs:
                                    run.font.name = '宋体'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体'); run.font.size = Pt(10)

                    merged_doc.add_paragraph("\n支付银行账户信息如下：")
                    binfo = [
                        ("名称", bank_info.get("名称", "未配置")),
                        ("银行类别", bank_info.get("银行类别", "未配置")),
                        ("开户银行名称", bank_info.get("开户银行名称", "未配置")),
                        ("银行账号", bank_info.get("银行账号", "未配置"))
                    ]
                    for k, v in binfo:
                        p_bank = merged_doc.add_paragraph(f"{k}：{v}" if k != "名称" else f"{k}，{v}")
                        for run in p_bank.runs:
                            run.font.name = '宋体'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                    merged_doc.add_paragraph("\n")
                    sig_name = sys_settings.get("company_signature", "省公众人力部")
                    p_sig = merged_doc.add_paragraph(sig_name)
                    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p_date = merged_doc.add_paragraph(datetime.datetime.now().strftime("%Y年%m月%d日"))
                    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for p in [p_sig, p_date]:
                        for run in p.runs:
                            run.font.name = '宋体'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体'); run.font.size = Pt(12)

            # [终极安全输出] 生成并抛出唯一的 Word 合集大包
            doc_io = io.BytesIO()
            merged_doc.save(doc_io)
            st.download_button(
                f"📥 点击下载【合并打印版】结算函 ({s_m}至{e_m})",
                data=doc_io.getvalue(),
                file_name=f"全量地市结算函_合并打印版_{s_m}至{e_m}.docx",
                type="primary"
            )

# ------------------------------------------------------------------------------
# Tab 3: 全局规则与参数配置
# ------------------------------------------------------------------------------
with tab3:
    st.subheader("🛠️ 全局算力引擎参数设置")
    st.info("💡 此处的设置将直接决定全公司数千条社保数据的最终核算金额。")

    c_y, c_e = st.columns(2)
    with c_y: target_year = st.selectbox("📅 规则生效年度", ["2024", "2025", "2026", "2027", "2028"], index=1)
    with c_e: target_entity = st.selectbox("🏢 适用法人主体", ["全量设置", "省公众", "中电数智", "省公司"])

    fetch_entity = "省公众" if "全量" in target_entity else target_entity
    curr = get_policy_rules(target_year, fetch_entity)

    if curr: st.success(f"已加载 {target_year} 年度 【{fetch_entity}】 的历史配置，可直接修改。")
    else: st.info(f"【{fetch_entity}】 在 {target_year} 年度尚无配置记录，请录入参数后保存。")

    if st.session_state.get('show_confirm', False):
        st.warning("⚠️ 警告：您即将修改底层财务参数，请核对变更情况！")
        if st.button("🔥 确认无误，强行写入数据库", type="primary", key="confirm_upsert_btn"):
            params = st.session_state.get('pending_params')
            if params:
                success, msg = upsert_policy_rules(params, is_all_entities=("全量" in params[1]))
                if success:
                    st.success(msg)
                    st.session_state['show_confirm'] = False
                    st.rerun()
                else: st.error(f"🚨 数据库写入被底层拦截！真实原因: {msg}")
        if st.button("❌ 撤销修改，返回编辑", key="cancel_upsert_btn"):
            st.session_state['show_confirm'] = False
            st.rerun()
        st.divider()

    with st.form("policy_rules_form"):
        st.write(f"**【{target_year} 年度】算法控制总开关**")
        c_mode, c_fund, c_med = st.columns(3)
        with c_mode:
            r_keys = ['exact', 'round_to_yuan', 'round_to_ten', 'floor_to_ten']
            cur_round = curr.get('rounding_mode', 'round_to_yuan')
            sel_round = st.selectbox("社保取整规则", options=r_keys, index=r_keys.index(cur_round) if cur_round in r_keys else 1)
        with c_fund:
            f_keys = ['independent', 'reverse_from_ss']
            cur_fund = curr.get('fund_calc_method', 'reverse_from_ss')
            sel_fund = st.selectbox("公积金特殊算法", options=f_keys, index=f_keys.index(cur_fund) if cur_fund in f_keys else 1)
        with c_med:
            med_serious = st.number_input("大病医疗个人固定扣款", value=safe_float(curr.get('medical_serious_fix', 7.0)), step=1.0)

        st.divider()
        def render_ins_row(label, prefix, has_pers=True, has_limit=True):
            cols = st.columns([1.5, 2, 2, 2, 2])
            cols[0].markdown(f"**{label}**")
            up = cols[1].number_input(f"{label}封顶", value=safe_float(curr.get(f'{prefix}_upper')), step=100.0, label_visibility="collapsed") if has_limit else 0.0
            lw = cols[2].number_input(f"{label}保底", value=safe_float(curr.get(f'{prefix}_lower')), step=100.0, label_visibility="collapsed") if has_limit else 0.0
            cr = cols[3].number_input(f"{label}企%", value=safe_float(curr.get(f'{prefix}_comp_rate')), step=0.01, format="%.4f", label_visibility="collapsed")
            pr = cols[4].number_input(f"{label}个%", value=safe_float(curr.get(f'{prefix}_pers_rate')), step=0.01, format="%.4f", label_visibility="collapsed") if has_pers else 0.0
            return up, lw, cr, pr

        p_up, p_lw, p_cr, p_pr = render_ins_row("养老保险", "pension")
        m_up, m_lw, m_cr, m_pr = render_ins_row("医疗保险", "medical")
        u_up, u_lw, u_cr, u_pr = render_ins_row("失业保险", "unemp")
        i_up, i_lw, i_cr, _ = render_ins_row("工伤保险", "injury", has_pers=False)
        mat_up, mat_lw, mat_cr, _ = render_ins_row("生育保险", "maternity", has_pers=False)
        f_up, f_lw, f_cr, f_pr = render_ins_row("住房公积金(官方线)", "fund")

        cols_soe = st.columns([1.5, 2, 2, 2, 2])
        cols_soe[0].markdown("**↳ 内部实际执行线**")
        f_soe_up = cols_soe[1].number_input("内部封顶", value=safe_float(curr.get('fund_soe_upper')), step=100.0, label_visibility="collapsed")
        f_soe_lw = cols_soe[2].number_input("内部保底", value=safe_float(curr.get('fund_soe_lower')), step=100.0, label_visibility="collapsed")

        _, _, a_cr, a_pr = render_ins_row("企业年金", "annuity", has_limit=False)

        if st.form_submit_button("🔍 对比并预览修改", type="primary"):
            st.session_state['pending_params'] = (
                target_year, target_entity, sel_round, sel_fund, med_serious,
                p_up, p_lw, p_cr, p_pr, m_up, m_lw, m_cr, m_pr,
                u_up, u_lw, u_cr, u_pr, i_up, i_lw, i_cr,
                mat_up, mat_lw, mat_cr, f_up, f_lw, f_cr, f_pr,
                a_cr, a_pr, f_soe_up, f_soe_lw
            )
            st.session_state['show_confirm'] = True
            st.rerun()

# ------------------------------------------------------------------------------
# Tab 4: 历史账单导入 (冷启动专区)
# ------------------------------------------------------------------------------
with tab4:
    st.subheader("📥 历史月度死账导入引擎 (冷启动与时光机)")
    st.warning("⚠️ 极度危险操作！此功能用于系统上线前的历史账单补录。利用“时光机”功能，可自动抓取最近月份的数据作为底稿，免去手动查工号的痛苦。")

    conn = _get_db_connection()
    try:
        archived_months_df = pd.read_sql_query("SELECT DISTINCT cost_month FROM ss_monthly_records ORDER BY cost_month DESC", conn)
        archived_months = archived_months_df['cost_month'].tolist()
    except Exception:
        archived_months = []
    finally:
        conn.close()

    hc1, hc2 = st.columns(2)
    with hc1:
        target_hist_month = st.text_input("📅 1. 设定你要补录的目标月份 (如: 2026-02)", value="2026-02")

        # [核心逻辑重构：智能寻轨] 计算绝对时间差，寻找“最近”的月份作为底稿母本
        ref_month = None
        if archived_months:
            try:
                def ym_to_int(ym_str):
                    y, m = map(int, str(ym_str).split('-'))
                    return y * 12 + m
                target_val = ym_to_int(target_hist_month)
                ref_month = min(archived_months, key=lambda x: abs(ym_to_int(x) - target_val))
            except Exception:
                ref_month = archived_months[0]

        if st.button("🚀 2. 生成带姓名与金额的【智能预填底稿】", type="primary"):
            conn = _get_db_connection()
            if ref_month:
                st.info(f"💡 探针已激活：系统侦测到最近的关联账单为 {ref_month}，正在为您完美克隆该月数据作为底稿...")
                clone_query = """
                    SELECT r.*, e.name AS '姓名'
                    FROM ss_monthly_records r
                    LEFT JOIN employees e ON r.emp_id = e.emp_id
                    WHERE r.cost_month = ?
                """
                ref_df = pd.read_sql_query(clone_query, conn, params=[ref_month])
            else:
                st.info("💡 探针已激活：系统暂无历史账单，正在从现有人事档案中提取名单生成空白底稿...")
                clone_query = """
                    SELECT e.emp_id, e.name AS '姓名', IFNULL(m.cost_center, '本级') AS cost_center
                    FROM employees e
                    LEFT JOIN ss_emp_matrix m ON e.emp_id = m.emp_id
                    WHERE e.status IN ('在职', '挂靠人员')
                """
                ref_df = pd.read_sql_query(clone_query, conn)
            conn.close()

            hist_cols = [
                '核算月份(YYYY-MM)', '工号', '姓名', '财务归属(成本中心)',
                '养老_企金额', '养老_个金额', '养老_通道(如:省公众)',
                '医疗_企金额', '医疗_个金额', '大病_个固定', '医疗_通道(如:省公司)',
                '失业_企金额', '失业_个金额', '失业_通道(如:省公众)',
                '工伤_企金额', '工伤_通道(如:省公众)',
                '生育_企金额', '生育_通道(如:省公司)',
                '公积金_企金额', '公积金_个金额', '公积金_通道(如:省公众)',
                '年金_企金额', '年金_个金额', '年金_通道(如:省公司)'
            ]

            output_data = []
            for _, row in ref_df.iterrows():
                out_row = {col: 0.0 for col in hist_cols}
                out_row['核算月份(YYYY-MM)'] = target_hist_month
                out_row['工号'] = row.get('emp_id', '')
                out_row['姓名'] = row.get('姓名', '')
                out_row['财务归属(成本中心)'] = row.get('cost_center', '本级')

                if ref_month:
                    out_row['养老_企金额'] = row.get('pension_comp', 0.0)
                    out_row['养老_个金额'] = row.get('pension_pers', 0.0)
                    out_row['养老_通道(如:省公众)'] = row.get('pension_route', '')
                    out_row['医疗_企金额'] = row.get('medical_comp', 0.0)
                    out_row['医疗_个金额'] = row.get('medical_pers', 0.0)
                    out_row['大病_个固定'] = row.get('medical_serious_pers', 0.0)
                    out_row['医疗_通道(如:省公司)'] = row.get('medical_route', '')
                    out_row['失业_企金额'] = row.get('unemp_comp', 0.0)
                    out_row['失业_个金额'] = row.get('unemp_pers', 0.0)
                    out_row['失业_通道(如:省公众)'] = row.get('unemp_route', '')
                    out_row['工伤_企金额'] = row.get('injury_comp', 0.0)
                    out_row['工伤_通道(如:省公众)'] = row.get('injury_route', '')
                    out_row['生育_企金额'] = row.get('maternity_comp', 0.0)
                    out_row['生育_通道(如:省公司)'] = row.get('maternity_route', '')
                    out_row['公积金_企金额'] = row.get('fund_comp', 0.0)
                    out_row['公积金_个金额'] = row.get('fund_pers', 0.0)
                    out_row['公积金_通道(如:省公众)'] = row.get('fund_route', '')
                    out_row['年金_企金额'] = row.get('annuity_comp', 0.0)
                    out_row['年金_个金额'] = row.get('annuity_pers', 0.0)
                    out_row['年金_通道(如:省公司)'] = row.get('annuity_route', '')

                output_data.append(out_row)

            hist_template = pd.DataFrame(output_data)
            hist_buffer = io.BytesIO()
            with pd.ExcelWriter(hist_buffer, engine='xlsxwriter') as writer:
                hist_template.to_excel(writer, index=False)
                writer.sheets['Sheet1'].set_column('A:X', 18)

            st.download_button("📥 3. 下载已预填好的历史账单底稿", data=hist_buffer.getvalue(), file_name=f"{target_hist_month}_历史社保智能补录底稿.xlsx", type="secondary")

    with hc2:
        hist_file = st.file_uploader("📤 4. 上传核对完毕的历史账单明细", type=["xlsx", "csv"])
        if hist_file and st.button("🚨 5. 强行覆写历史月度账单入库", type="primary"):
            try:
                h_df = pd.read_csv(hist_file) if hist_file.name.endswith('.csv') else pd.read_excel(hist_file)
                conn = _get_db_connection()
                cursor = conn.cursor()

                upsert_sql = """
                    INSERT INTO ss_monthly_records (
                        record_id, cost_month, emp_id, cost_center,
                        pension_comp, pension_pers, pension_route,
                        medical_comp, medical_pers, medical_serious_pers, medical_route,
                        unemp_comp, unemp_pers, unemp_route,
                        injury_comp, injury_route,
                        maternity_comp, maternity_route,
                        fund_comp, fund_pers, fund_route,
                        annuity_comp, annuity_pers, annuity_route
                    ) VALUES (?,?,?,?, ?,?,?, ?,?,?,?, ?,?,?, ?,?, ?,?, ?,?,?, ?,?,?)
                    ON CONFLICT(cost_month, emp_id) DO UPDATE SET
                        cost_center=excluded.cost_center,
                        pension_comp=excluded.pension_comp, pension_pers=excluded.pension_pers, pension_route=excluded.pension_route,
                        medical_comp=excluded.medical_comp, medical_pers=excluded.medical_pers, medical_serious_pers=excluded.medical_serious_pers, medical_route=excluded.medical_route,
                        unemp_comp=excluded.unemp_comp, unemp_pers=excluded.unemp_pers, unemp_route=excluded.unemp_route,
                        injury_comp=excluded.injury_comp, injury_route=excluded.injury_route,
                        maternity_comp=excluded.maternity_comp, maternity_route=excluded.maternity_route,
                        fund_comp=excluded.fund_comp, fund_pers=excluded.fund_pers, fund_route=excluded.fund_route,
                        annuity_comp=excluded.annuity_comp, annuity_pers=excluded.annuity_pers, annuity_route=excluded.annuity_route
                """
                count = 0
                for _, row in h_df.iterrows():
                    h_month = str(row.get('核算月份(YYYY-MM)', '')).strip()
                    eid = str(row.get('工号', '')).replace('.0', '').strip()
                    if not h_month or not eid or h_month == 'nan' or eid == 'nan': continue

                    rec_id = f"{h_month}_{eid}"
                    cc = str(row.get('财务归属(成本中心)', '本级')).strip()

                    cursor.execute(upsert_sql, (
                        rec_id, h_month, eid, cc,
                        safe_float(row.get('养老_企金额')), safe_float(row.get('养老_个金额')), str(row.get('养老_通道(如:省公众)', 'None')),
                        safe_float(row.get('医疗_企金额')), safe_float(row.get('医疗_个金额')), safe_float(row.get('大病_个固定')), str(row.get('医疗_通道(如:省公司)', 'None')),
                        safe_float(row.get('失业_企金额')), safe_float(row.get('失业_个金额')), str(row.get('失业_通道(如:省公众)', 'None')),
                        safe_float(row.get('工伤_企金额')), str(row.get('工伤_通道(如:省公众)', 'None')),
                        safe_float(row.get('生育_企金额')), str(row.get('生育_通道(如:省公司)', 'None')),
                        safe_float(row.get('公积金_企金额')), safe_float(row.get('公积金_个金额')), str(row.get('公积金_通道(如:省公众)', 'None')),
                        safe_float(row.get('年金_企金额')), safe_float(row.get('年金_个金额')), str(row.get('年金_通道(如:省公司)', 'None'))
                    ))
                    count += 1
                conn.commit()
                st.success(f"✅ 历史死账冷启动成功！共覆盖/写入 {count} 条月度固化记录。")
            except Exception as e:
                st.error(f"❌ 导入崩溃: {e}")
            finally:
                if 'conn' in locals(): conn.close()